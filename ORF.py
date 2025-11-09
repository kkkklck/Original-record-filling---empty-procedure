
# === 原始记录自动填写程序 ===


from pathlib import Path
import re, copy, math, warnings, sys, os, unicodedata, ctypes
from collections import defaultdict
from typing import Union
from docx import Document
from docx.shared import RGBColor, Pt
from openpyxl.styles import Font, Alignment


warnings.filterwarnings("ignore", category=SyntaxWarning)

TITLE = "原始记录自动填写程序"
VERSION = "v 1.0.1"


# ===== 默认路径 =====
SCRIPT_DIR = Path(__file__).resolve().parent
WORD_SRC_DEFAULT = Path(r"D:\eg\eg.docx")
XLSX_WITH_SUPPORT_DEFAULT = SCRIPT_DIR / "原始记录excel模板.xlsx"
DEFAULT_FONT_PT = 9

# 每页 5 组、每组 5 行、每行 8 读数+平均值
PER_LINE_PER_BLOCK = 5
BLOCKS_PER_SHEET = 5
MU_DIGITS_THRESHOLD = 4  # 需求：四位数→μ

# 本次运行只提示一次
_hint_shown = False

# 打印顺序：可自行调整位置
CATEGORY_ORDER = ["钢柱", "钢梁", "支撑", "网架", "其他"]

# 支撑/网架 分桶策略："number"=按编号，"floor"=按楼层；仅本次运行生效
support_bucket_strategy = None
net_bucket_strategy = None

# 轻量识别缓存：避免重复读取 Word
_PROBE_CACHE = {
    "src": None,
    "grouped": None,
    "all_rows": None,
    "categories": None,
}

# —— 严防跨类/跨 μ 写串（开关）——
STRICT_CROSS_CAT_GUARD = True

def _sheet_cat_from_title(title: str) -> str | None:
    """根据 sheet 名推断类别：去掉（n）和 μ 后比对前缀。"""
    base = re.sub(r"（\d+）$", "", (title or "").strip())
    base = base.replace(" μ", "μ")  # 容错：有人手抖加空格
    base_wo_mu = base.replace("μ", "")
    for c in CATEGORY_ORDER:
        if base_wo_mu.startswith(c):
            return c
    return None

def _is_mu_title(title: str) -> bool:
    return "μ" in (title or "")

def _filter_pages_for_cat(pages: list[str], cat: str) -> list[str]:
    """只保留属于 cat 的页名（再保险）。"""
    return [p for p in pages if _sheet_cat_from_title(p) == cat]


# === 通用输入封装 ===
def enable_ansi():
    if os.name != "nt":
        return True
    k32 = ctypes.windll.kernel32
    h = k32.GetStdHandle(-11)  # STD_OUTPUT_HANDLE
    mode = ctypes.c_uint32()
    if not k32.GetConsoleMode(h, ctypes.byref(mode)):
        return False
    return bool(k32.SetConsoleMode(h, mode.value | 0x0004))  # ENABLE_VIRTUAL_TERMINAL_PROCESSING


enable_ansi()







class BackStep(Exception):
    """用户输入 q 请求返回上一步。"""
    pass


class AbortToPath(Exception):
    """用户主动中断当前模式并返回路径输入。"""
    pass


def ask(prompt: str, allow_empty: bool = True, lower: bool = False) -> str:
    """统一的控制台输入函数。

    参数:
        prompt: 提示字符串。
        allow_empty: 是否允许空输入；False 时会重复询问。
        lower: 返回值是否小写化。

    返回:
        用户输入的字符串（可小写化）。

    特殊:
        输入 ``q`` 将触发 :class:`BackStep` 异常。
        仅识别小写 ``q``，大写 ``Q`` 在此阶段视为普通字符。
    """
    while True:
        raw = input(f"{prompt}\n→ ").strip()
        if raw == "q":
            raise BackStep()
        if not allow_empty and raw == "":
            continue
        return raw.lower() if lower else raw


def show_help_browser():
    """帮助浏览器包装。"""
    tutorial_browser()



def ask_path() -> str | None:
    """顶层路径输入。

    返回 ``None`` 表示用户查看帮助后继续；
    返回 ``"__QUIT__"`` 表示用户请求退出程序；
    其他返回值为用户输入的路径字符串。
    """
    raw = input("📂 请输入 Word 源路径（eg：D:\示例.docx）\n→ ").strip()
    if raw == "help":
        show_help_browser()
        return None
    if raw == "Q":
        return "__QUIT__"
    return raw


def is_valid_path(p: str) -> bool:
    """简单校验路径是否存在。"""
    path_obj = Path(p.strip('"'))
    return path_obj.exists() and path_obj.is_file()


# ---- 文件占用友好提示封装 ----
class FileInUse(Exception):
    pass


def _is_in_use_error(e: Exception) -> bool:
    # Windows 常见：WinError 32（共享冲突），或 PermissionError 13
    msg = str(e).lower()
    code32 = getattr(e, "winerror", None) == 32
    perm13 = isinstance(e, PermissionError)
    hit_msg = ("being used by another process" in msg or
               "used by another process" in msg or
               "permission denied" in msg)
    return bool(code32 or perm13 or hit_msg)


def load_workbook_safe(path, **kw):
    from openpyxl import load_workbook
    try:
        return load_workbook(path, **kw)
    except Exception as e:
        if _is_in_use_error(e):
            raise FileInUse(f"Excel 模板/文件被占用：{path}") from e
        raise


def save_workbook_safe(wb, path):
    try:
        wb.save(path)
    except Exception as e:
        if _is_in_use_error(e):
            raise FileInUse(f"无法保存 Excel（被占用）：{path}") from e
        raise


def save_docx_safe(doc, path):
    try:
        doc.save(str(path))
    except Exception as e:
        if _is_in_use_error(e):
            raise FileInUse(f"无法保存 Word（被占用）：{path}") from e
        raise


# ===== Word 汇总生成 =====
NEED_COLS = 11
MIN_ROWS_EACH = 5
PLACEHOLDER = "/"
digit_re = re.compile(r"\d")
HEADER = [
    "序号", "构件名称及部位",
    "测点1 读数1", "测点1 读数2",
    "测点2 读数1", "测点2 读数2",
    "测点3 读数1", "测点3 读数2",
    "测点4 读数1", "测点4 读数2",
    "涂层厚度平均值"
]


def ensure_cells(row, need=NEED_COLS):
    """
    确保表格行包含足够的单元格，不足时自动补充空白单元格。

    通过复制首个单元格的格式创建空白单元格，避免因原始表格列数不足导致数据提取失败，保障数据结构完整性。

    Args:
        row: Word表格行对象（docx.table.Row）
        need: 需要的最小列数，默认11列（与汇总表列数一致）
    """
    while len(row.cells) < need:
        tc = copy.deepcopy(row.cells[0]._tc)  # noqa
        for t in tc.xpath('.//*[local-name()="t"]'): t.text = ''
        row._tr.append(tc)  # noqa


def color_row_red(row):
    """
    将表格行的文字颜色设置为红色，用于表头高亮显示。

    通过遍历行内所有单元格和段落，统一设置文字颜色为红色，增强汇总表中表头与数据行的区分度。

    Args:
        row: Word表格行对象（docx.table.Row）
    """
    for c in row.cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)


def is_data_table(tbl):
    """
    判断Word表格是否为有效数据表格（含测点和平均值信息）。

    通过检查表格前3行是否同时包含“测点1”和“平均值”关键词，筛选出实际存储检测数据的表格，排除说明性表格。

    Args:
        tbl: Word表格对象（docx.table.Table）
    Returns:
        bool: 是有效数据表格则返回True，否则返回False
    """
    first_three = " ".join(c.text for r in tbl.rows[:3] for c in r.cells)
    return "测点1" in first_three and "平均值" in first_three


def detect_layout(tbl):
    """
    检测数据表格的列布局，确定测点列、平均值列位置及是否为钢梁表格。

    定位含“测点1”的表头行，提取测点列索引和平均值列索引；通过测点列数量判断是否为钢梁表格（钢梁含3个测点）。

    Args:
        tbl: Word表格对象（docx.table.Table）
    Returns:
        tuple: 包含三个元素的元组，分别为：
            - 测点列索引列表（list[int]）
            - 平均值列索引（int）
            - 是否为钢梁表格（bool，钢梁表格返回True）
    """
    hdr = next(r for r in tbl.rows if "测点1" in "".join(c.text for c in r.cells))
    col_vals, col_avg = [], None
    for i, t in enumerate(hdr.cells):
        txt = (t.text or "").strip()
        m = re.match(r"测点(\d+)", txt)
        if m:
            col_vals.append(i)
        elif "平均值" in txt and "所有" not in txt:
            col_avg = i
    is_beam = len(col_vals) == 3  # 梁 3 组，柱/支撑 4 组
    return col_vals, col_avg, is_beam


def extract_rows_with_progress(tbl, ti: int, T: int, *, show_progress: bool = True):  # noqa
    """
    从数据表格提取行数据，带实时进度提示。

    按表头布局提取构件名称、测点值和平均值，对钢梁表格自动补充第4个测点（用“/”占位）；通过控制台实时显示提取进度（按行计算）。

    Args:
        tbl: Word表格对象（docx.table.Table）
        ti: 当前表格在总表格中的序号（从1开始）
        T: 需处理的总表格数量
    Returns:
        list[dict]: 提取的数据行列表，每个元素为包含以下键的字典：
            - name: 构件名称（str）
            - vals: 测点值列表（list[str]）
            - avg: 平均值（str）
            - is_hdr: 是否为表头行（bool）
    """
    col_vals, col_avg, is_beam = detect_layout(tbl)
    rows, last_comp, last_avg = [], None, ""
    buffer = []

    total = len(tbl.rows)
    last_flush = -1

    for ridx, r in enumerate(tbl.rows):
        if show_progress and ridx // 20 != last_flush:
            last_flush = ridx // 20
            pct = int((ridx + 1) * 100 / max(1, total))
            sys.stdout.write(f"\r📝 读取 Word：表 {ti}/{T}（{pct}%）")
            sys.stdout.flush()

        line = " ".join(c.text for c in r.cells)

        if "测点1" in line:
            if buffer:
                rows.extend(buffer);
                buffer.clear()  # noqa
            meas_titles = [f"测点{i + 1}" for i in range(len(col_vals))]
            if is_beam: meas_titles.append("测点4")  # 梁补第4列标题
            rows.append({"name": "", "vals": meas_titles, "avg": "平均值", "is_hdr": True})
            continue

        if not digit_re.search(line):
            continue

        comp = r.cells[1].text.strip()
        vals = [r.cells[i].text.strip() for i in col_vals]
        if is_beam and len(vals) == 3: vals.append("/")

        raw_avg = r.cells[col_avg].text.replace("\n", "").strip()
        avg = raw_avg or last_avg or "/"
        last_avg = avg if raw_avg else last_avg

        buffer.append({"name": comp if comp != last_comp else "",
                       "vals": vals, "avg": avg, "is_hdr": False})
        last_comp = comp

    rows.extend(buffer)
    sys.stdout.write(f"\r📝 读取 Word：表 {ti}/{T}（100%）\n");
    sys.stdout.flush()
    return rows


def build_summary_doc_with_progress(rows):
    """
     生成Word汇总表，带实时进度提示。

     将提取的数据行整理为规范表格，表头标红；不足行数用占位符补充，统一字体大小；通过控制台显示组装进度。

     Args:
         rows: 提取的数据行列表（extract_rows_with_progress返回结果）
     Returns:
         Document: 生成的Word汇总表文档对象（docx.document.Document）
     """
    doc = Document()
    tbl = doc.add_table(rows=1, cols=NEED_COLS)
    tbl.style = "Table Grid"
    for i, t in enumerate(HEADER):
        tbl.rows[0].cells[i].text = t
    color_row_red(tbl.rows[0])

    serial, last_comp, buffer = 1, None, []
    total = len(rows)
    step = max(50, total // 100)

    def flush():
        nonlocal serial, buffer
        miss = max(0, MIN_ROWS_EACH - len(buffer))
        for _ in range(miss):
            q = tbl.add_row();
            ensure_cells(q)
            for z in range(2, 10): q.cells[z].text = PLACEHOLDER
            q.cells[10].text = PLACEHOLDER
        serial += 1;
        buffer.clear()

    for i, it in enumerate(rows, start=1):
        if i % step == 0 or i == total:
            pct = int(i * 100 / max(1, total))
            sys.stdout.write(f"\r📦 组装汇总：{i}/{total}（{pct}%）")
            sys.stdout.flush()

        if it["is_hdr"] and buffer: flush()

        raw_name = (it.get("name") or "").strip()
        comp = raw_name or last_comp or ""

        if last_comp and comp and comp != last_comp:
            flush();
            last_comp = None

        if it.get("is_hdr"):
            r = tbl.add_row();
            ensure_cells(r);
            color_row_red(r)
            r.cells[1].text = "构件名称及部位" if not raw_name else raw_name
            for k, v in enumerate(it["vals"]):
                c = 2 + k * 2
                r.cells[c].text = v
            r.cells[10].text = it["avg"]
            last_comp = comp
            continue

        r = tbl.add_row();
        ensure_cells(r);
        buffer.append(r)
        first = (last_comp is None) or (comp and comp != last_comp)
        if first:
            r.cells[0].text = str(serial)
            r.cells[1].text = raw_name
            last_comp = comp
        for k, v in enumerate(it["vals"]):
            c = 2 + k * 2
            r.cells[c].text = v
            r.cells[c + 1].text = v
        r.cells[10].text = it["avg"]

    flush()
    sys.stdout.write("\n");
    sys.stdout.flush()
    return doc


def set_doc_font_progress(doc, pt=DEFAULT_FONT_PT):
    """
    统一Word文档中所有文字的字体大小，带实时进度提示。

    遍历文档中的所有段落和表格单元格，将字体大小设置为指定磅数（默认9pt）；通过控制台显示字体设置进度。

    Args:
        doc: Word文档对象（docx.document.Document）
        pt: 字体大小（磅），默认9pt
    """
    cell_pars = 0
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                cell_pars += len(c.paragraphs)
    total = len(doc.paragraphs) + cell_pars
    done = 0
    step = max(200, total // 100)

    for p in doc.paragraphs:
        for run in p.runs: run.font.size = Pt(pt)
        done += 1
        if done % step == 0 or done == total:
            pct = int(done * 100 / max(1, total))
            sys.stdout.write(f"\r🖋 统一字体：{done}/{total}（{pct}%）");
            sys.stdout.flush()

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs: run.font.size = Pt(pt)
                    done += 1
                    if done % step == 0 or done == total:
                        pct = int(done * 100 / max(1, total))
                        sys.stdout.write(f"\r🖋 统一字体：{done}/{total}（{pct}%）");
                        sys.stdout.flush()
    sys.stdout.write("\n");
    sys.stdout.flush()


# ===== rows → groups（8读数+平均值）=====
def groups_from_your_rows(rows_all_tables):
    """
    将提取的原始数据行转换为按构件分组的结构化数据。

    按构件名称分组，将每组数据整理为规范格式（8个读数+1个平均值），自动用“/”补齐不足的读数。

    Args:
        rows_all_tables: 所有表格提取的原始数据行列表（extract_rows_with_progress返回结果）
    Returns:
        list[dict]: 构件数据组列表，每个元素为包含以下键的字典：
            - name: 构件名称（str）
            - data: 数据行列表，每行包含8个读数和1个平均值（list[list[str]]）
    """
    groups = [];
    cur = None
    for it in rows_all_tables:
        if it.get("is_hdr"): continue
        name = (it.get("name") or "").strip()
        if name:
            if cur and cur["data"]: groups.append(cur)  # noqa
            cur = {"name": name, "data": []}
        if not cur: continue
        vals8 = []
        for v in it["vals"]:
            v = (v or "/").strip() or "/"
            vals8.extend([v, v])
        while len(vals8) < 8: vals8.append("/")
        avg = (it.get("avg") or "/").strip() or "/"
        cur["data"].append(vals8[:8] + [avg])  # noqa
    if cur and cur["data"]: groups.append(cur)
    return groups


# ===== 分类 / 规则 =====
CATEGORY_SYNONYMS = {
    "网架": [
        "网架", "WJ", "SPACE FRAME", "SPACEFRAME", "GRID", "GRID STRUCTURE",
        "桁架网架", "球节点", "网壳", "SJ",
        "XX", "SX", "FG", "上弦", "下弦", "腹杆"
    ],
    "支撑": ["支撑", "WZ", "ZC", "支架", "斜撑", "撑杆"],
    "钢柱": ["钢柱", "柱", "GZ", "框架柱", "立柱", "H柱"],
    "钢梁": ["钢梁", "梁", "GL", "连系梁", "檩条", "楼梯梁", "平台梁", "屋架梁"],
}


def kind_of(name: str) -> str:
    """
    根据构件名称判断类型（钢柱/钢梁/支撑/其他）。

    基于预设的同义词表匹配构件名称中的关键词（如“钢柱”或“GZ”对应钢柱），未匹配到关键词的构件归为“其他”类。

    Args:
        name: 构件名称字符串（str）
    Returns:
        str: 构件类型，可能为"钢柱"、"钢梁"、"支撑"或"其他"
    """
    s_up = name.upper()
    for cat, words in CATEGORY_SYNONYMS.items():
        for w in words:
            if w.isascii():
                if w.upper() in s_up:
                    return cat
            else:
                if w in name:
                    return cat
    return "其他"  # 未识别 → 其他


def floor_of(name: str) -> int:
    """
    从构件名称中提取楼层号，特殊楼层用固定大数值标记。
    规则更新：
      - 机房优先于屋面，避免“屋面机房层”被误判为屋面
      - 机房层: 10**6 - 1；屋面: 10**6
    """
    s = (name or "").replace("－", "-").replace("—", "-").replace("–", "-")
    sl = s.lower()
    # ① 先机房（更具体）
    if re.search(r"(机房(?:层)?|\bjf\b)", sl):
        return 10 ** 6 - 1
    # ② 再屋面
    if re.search(r"(?:屋面|屋顶|顶\s*层)", s) or re.search(r"\b(?:wm|dc)\b", sl):
        return 10 ** 6
    # ③ 常规数字层
    m = re.search(r"(?i)[FL]\s*(\d+)", s)
    if m: return int(m.group(1))
    m = re.search(r"(?i)(\d+)\s*[FL]", s)
    if m: return int(m.group(1))
    m = re.search(r"(\d+)\s*[层樓楼]", s)
    if m: return int(m.group(1))
    # ④ 地下/负层 → 统归 0（排序靠 _floor_label_from_name）
    if re.search(r"(?i)\bB\s*\d+\b|负\s*\d+\s*层?", s):
        return 0
    return 0



def _floor_label_from_name(name: str) -> str:
    """返回标签：B2 / 5F / 机房层 / 屋面 ...（机房优先于屋面）"""
    s = (name or "").replace("－", "-").replace("—", "-").replace("–", "-")
    sl = s.lower()
    # ① 机房先判
    if re.search(r"(机房(?:层)?|\bjf\b)", sl):
        return "机房层"
    # ② 再屋面
    if re.search(r"屋面|顶层", s) or re.search(r"\b(?:wm|dc)\b", sl):
        return "屋面"
    m = re.search(r"(?i)B\s*(\d+)", s)
    if m: return f"B{int(m.group(1))}"
    m = re.search(r"(\d+)\s*[Ff层樓楼]?", s)
    if m: return f"{int(m.group(1))}F"
    return "F?"



def _floor_sort_key_by_label(label: str):
    """生成楼层标签的排序键。"""
    m = re.fullmatch(r"B(\d+)", label)
    if m:
        return (0, -int(m.group(1)))
    m = re.fullmatch(r"(\d+)F", label)
    if m:
        return (1, int(m.group(1)))
    if label == "机房层":
        return (2, 0)
    if label == "屋面":
        return (3, 0)
    return (4, 0)


def segment_index(floor: int, breaks: list[int]) -> int:
    """
    根据楼层断点返回分段索引。
    更新：
      - 机房层单独成段，位于数字层之后、屋面之前
      - 屋面在最末段
      - 若无断点，机房层与屋面也能稳定分开（索引 1 / 2）
    """
    # 没有断点时：0=数字&地下，1=机房层，2=屋面
    if not breaks:
        if floor == 10**6 - 1:  # 机房层
            return 1
        if floor >= 10**6:      # 屋面
            return 2
        return 0

    # 有断点时：数字层 → 机房层( len(breaks) ) → 屋面( len(breaks)+1 )
    if floor == 10**6 - 1:      # 机房层
        return len(breaks)
    if floor >= 10**6:          # 屋面
        return len(breaks) + 1

    # 常规数字层：落到第一个 >= 断点 的段
    for i, b in enumerate(breaks):
        if floor <= b:
            return i
    return len(breaks)  # 高于最大断点的数字层（极高层）仍落在最后一个数字段



def expand_blocks(groups, block_size=PER_LINE_PER_BLOCK):
    """
    将构件数据组拆分为固定大小的数据块（默认5行/块），不足行数用“/”补齐。

    按指定块大小（默认5行）拆分每组数据，确保每个块结构统一，适配Excel模板中“每组数据占5行”的格式要求。

    Args:
        groups: 构件数据组列表（groups_from_your_rows返回结果）
        block_size: 每个数据块的行数，默认5行
    Returns:
        list[dict]: 数据块列表，每个元素为包含以下键的字典：
            - name: 构件名称（str）
            - data: 5行数据（每行9列，list[list[str]]）
    """
    blocks = []
    for g in groups:
        rows = list(g["data"])
        for k in range(0, len(rows), block_size):
            sub = rows[k:k + block_size]
            while len(sub) < block_size: sub.append(['/'] * 9)
            blocks.append({"name": g["name"], "data": sub})
    return blocks


# ===== Excel sheet 复制与设置 =====
def clone_sheet_keep_print(wb, tpl_name: str, title: str):
    """
    复制Excel工作表并保留打印格式和视图设置，确保新表与模板格式一致。

    复制内容包括视图（缩放、冻结窗格）、打印区域、页面设置（方向、纸张大小）、页边距、行列宽等，保障格式统一性。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        tpl_name: 模板工作表名称（str）
        title: 新工作表名称（str）
    Returns:
        openpyxl.worksheet.worksheet.Worksheet: 新复制的工作表对象
    """
    tpl = wb[tpl_name]
    ws = wb.copy_worksheet(tpl)
    ws.title = title
    ws.sheet_view.view = "pageBreakPreview"
    try:
        ws.freeze_panes = tpl.freeze_panes
    except:
        pass
    try:
        ws.print_area = tpl.print_area
    except:
        pass
    try:
        ws.print_titles = tpl.print_titles
    except:
        pass
    for attr in (
            "orientation", "paperSize", "fitToWidth", "fitToHeight", "scale", "firstPageNumber", "useFirstPageNumber"):
        try:
            setattr(ws.page_setup, attr, getattr(tpl.page_setup, attr))
        except:
            pass
    for attr in ("left", "right", "top", "bottom", "header", "footer"):
        try:
            setattr(ws.page_margins, attr, getattr(tpl.page_margins, attr))
        except:
            pass
    for col, dim in tpl.column_dimensions.items():
        if dim.width is not None:
            ws.column_dimensions[col].width = dim.width
    for row, dim in tpl.row_dimensions.items():
        if dim.height is not None:
            ws.row_dimensions[row].height = dim.height
    return ws


def ensure_total_pages(wb, base: str, total_needed: int):
    """
    确保Excel中有足够的指定类型工作表，不足时自动从基础表复制补充。

    筛选并排序已有同类型工作表，若数量不足，以基础表为模板复制新表并按序号命名（如“钢柱（2）”）。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        base: 基础工作表名称（如"钢柱"，str）
        total_needed: 需要的工作表总数（int）
    Returns:
        list[str]: 排序后的工作表名称列表
    """
    names = [s for s in wb.sheetnames if s == base or re.match(rf'^{re.escape(base)}（\d+）$', s)]
    names = sorted(names, key=lambda n: 0 if n == base else int(re.findall(r'（(\d+)）', n)[0]))
    have = len(names)
    start = have + 1
    for _ in range(max(0, total_needed - have)):
        nm = f"{base}（{start}）"
        clone_sheet_keep_print(wb, base, nm)
        names.append(nm);
        start += 1
    return names


def ensure_total_pages_from(wb, tpl_name: str, new_base: str, total_needed: int):
    """
    为“其他”类构件确保足够的工作表，复用已有表或从指定模板复制。

    适用于无专用模板的类别，筛选已有同类型工作表，不足时从指定模板（如“钢柱”）复制新表并命名。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        tpl_name: 模板工作表名称（如"钢柱"，str）
        new_base: 新类别基础名称（如"其他"，str）
        total_needed: 需要的工作表总数（int）
    Returns:
        list[str]: 排序后的工作表名称列表
    """
    # 复用已有“其他（n）”等；不足则从 tpl_name 复制
    names = [s for s in wb.sheetnames if s == new_base or re.match(rf'^{re.escape(new_base)}（\d+）$', s)]
    names = sorted(names, key=lambda n: 0 if n == new_base else int(re.findall(r'（(\d+)）', n)[0]))
    have = len(names)
    start = have + 1
    for _ in range(max(0, total_needed - have)):
        nm = f"{new_base}（{start}）" if start > 1 else new_base
        clone_sheet_keep_print(wb, tpl_name, nm)
        if nm not in names:
            names.append(nm)
        start += 1
    return names

# ========= μ 判定 & μ 页创建 & 清理 =========

def _normalize_digits(s: str) -> str:
    """把全角/带逗号/空格/点的数字统一成 ASCII 连续数字串：'４,070.0' → '40700'。"""
    s = unicodedata.normalize("NFKC", str(s or ""))
    parts = re.findall(r"\d+", s)
    return "".join(parts)


def _is_mu_block(block: dict) -> bool:
    # """判断一个块是否含 μ 值。
    #
    # 仅检查每行前 8 个读数格：
    #   * 若某格为纯 \d{4,} 数字串 → μ；
    #   * 若某格能解析为数值且绝对值 ≥1000，且不含单位/文字 → μ。
    # """
    for row in block.get("data", []):
        cells = row[:8] if isinstance(row, (list, tuple)) else []
        for v in cells:
            if v in (None, "/", "／"):
                continue
            s = unicodedata.normalize("NFKC", str(v)).strip()
            if re.fullmatch(r"\d{4,}", s):
                return True
            if re.fullmatch(r"[\d.,]+", s):
                try:
                    if abs(float(s.replace(",", ""))) >= 1000:
                        return True
                except Exception:
                    pass
    return False


def _ensure_mu_pages_shared(wb, base: str, mu_tpl: str, start_idx: int, count: int) -> list[str]:
    """
    基于 μ 母版（如 '钢梁μ'）批量生成编号页，序号从 start_idx+1 起。
    若 start_idx 为 0，则复用模板页作为首张。
    返回生成（或已有）的 μ 页名列表：['钢梁μ', '钢梁μ（2）', ...]
    """
    pages = []
    use_tpl_first = start_idx == 0 and mu_tpl in wb.sheetnames
    for idx in range(start_idx + 1, start_idx + count + 1):
        if use_tpl_first and idx == start_idx + 1:
            pages.append(mu_tpl)
            continue
        nm = f"{base}μ（{idx}）"
        if nm not in wb.sheetnames:
            if mu_tpl not in wb.sheetnames:
                raise RuntimeError(f"缺少 μ 模板：{mu_tpl}")
            clone_sheet_keep_print(wb, mu_tpl, nm)
        pages.append(nm)
    return pages


def cleanup_unused_mu_templates(wb, used_pages: list[str]):
    """
    清掉本次没用到的“裸 μ 模板页”（如 '钢梁μ'）。
    """
    used = set(used_pages or [])
    base_candidates = ["钢柱μ", "钢梁μ", "支撑μ", "网架μ", "钢柱 μ", "钢梁 μ", "支撑 μ", "网架 μ"]
    for base in base_candidates:
        if base in wb.sheetnames and base not in used:
            try:
                wb.remove(wb[base])
            except Exception:
                pass
# ========= μ 分流 + 共用编号（通用分页器） =========
def split_mu_blocks(blocks):
    normal, mu = [], []
    for b in blocks:
        (mu if _is_mu_block(b) else normal).append(b)
    return normal, mu

def pages_needed(blocks):
    return math.ceil(len(blocks) / BLOCKS_PER_SHEET) if blocks else 0

def ensure_pages_slices_for_cat_muaware(wb, cat: str, blocks_by_bucket: dict[int, list]):
    """
    μ 逻辑的通用分页器（修正版）：
      - 同一桶内：先普通页、后 μ 页；同页不混
      - 序号共用：普通与 μ 跨桶连续编号
      - 只为需要的普通页创建 sheet，不会因为 μ 页而“补造”普通页
    返回：pages_slices、blocks_slices（按桶顺序的列表）
    """
    buckets = sorted(blocks_by_bucket.keys())
    pages_slices = []
    blocks_slices = []

    # 两套计数：一个用于“编号”（普通+μ），一个用于“普通页实际已创建数”
    total_all_pages = 0            # 普通 + μ，决定 μ 页的起始序号
    normal_pages_created = 0       # 仅普通页，决定 ensure_total_pages 的目标数

    for i in buckets:
        all_blocks = blocks_by_bucket.get(i, []) or []

        # 拆分普通/μ
        normal_blocks = []
        mu_blocks = []
        for b in all_blocks:
            (mu_blocks if _is_mu_block(b) else normal_blocks).append(b)

        need_n = math.ceil(len(normal_blocks) / BLOCKS_PER_SHEET) if normal_blocks else 0
        need_m = math.ceil(len(mu_blocks) / BLOCKS_PER_SHEET) if mu_blocks else 0

        # 1) 普通页：只按“普通页已创建数 + 本桶普通需求”来确保
        if need_n:
            normal_full = ensure_total_pages(wb, cat, normal_pages_created + need_n)
            # 取出“本桶新分配”的那一段
            normal_batch = normal_full[normal_pages_created : normal_pages_created + need_n]
            normal_pages_created += need_n
        else:
            normal_batch = []

        # 2) μ 页：编号要接在“已有总页数 + 本桶普通页数”之后
        #    但不需要为了编号去创建额外的“普通空页”
        mu_batch = []
        if need_m:
            start_idx_for_mu = total_all_pages + need_n  # 先算上同桶普通页
            mu_batch = _ensure_mu_pages_shared(
                wb, base=cat, mu_tpl=f"{cat}μ",
                start_idx=start_idx_for_mu, count=need_m
            )

        # 3) 更新“总页数”计数（普通+μ）
        total_all_pages += (need_n + need_m)

        # 4) 汇总本桶
        pages_slices.append(normal_batch + mu_batch)
        blocks_slices.append(normal_blocks + mu_blocks)

    return pages_slices, blocks_slices

# ===== 快速探测文档中包含的构件类别（供前端静默识别用） =====

def probe_categories_from_docx(src: Union[str, Path]) -> dict:
    """轻量识别 Word，返回类别顺序与数量，并写入缓存。"""
    p = Path(str(src)).resolve()
    if not p.exists():
        raise FileNotFoundError(f"未找到 Word 源文件：{p}")

    cache_src = _PROBE_CACHE.get("src")
    if cache_src and Path(str(cache_src)).resolve() == p:
        grouped_cached = _PROBE_CACHE.get("grouped") or {}
        cats_cached = _PROBE_CACHE.get("categories") or []
        counts_cached = {c: len(grouped_cached.get(c, [])) for c in cats_cached}
        for k in CATEGORY_ORDER:
            counts_cached.setdefault(k, 0)
        return {"categories": list(cats_cached), "counts": counts_cached}

    groups_all_tables, all_rows = read_groups_from_doc(p, progress=False)
    grouped = defaultdict(list)
    for g in groups_all_tables:
        grouped[kind_of(g["name"])].append(g)
    categories_present = [cat for cat in CATEGORY_ORDER if grouped.get(cat)]

    _PROBE_CACHE.update({
        "src": str(p),
        "grouped": grouped,
        "all_rows": all_rows,
        "categories": categories_present,
    })

    counts = {cat: len(grouped.get(cat, [])) for cat in categories_present}
    for k in CATEGORY_ORDER:
        counts.setdefault(k, 0)
    return {"categories": list(categories_present), "counts": counts}



def enforce_mu_font(wb):
    """
    遍历Excel所有单元格，将含“μ”字符的单元格字体强制设为Times New Roman。

    解决“μ”符号在部分字体下显示异常的问题，保留原字体的大小、加粗等其他属性。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
    """
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                v = cell.value
                if isinstance(v, str) and "μ" in v:
                    f = cell.font
                    cell.font = Font(
                        name="Times New Roman",
                        sz=f.sz, bold=f.bold, italic=f.italic, vertAlign=f.vertAlign,
                        underline=f.underline, strike=f.strike, color=f.color,
                        charset=f.charset, scheme=f.scheme, outline=f.outline
                    )


# ===== 数据区定位 / 写入 =====
def detect_anchors(ws):
    """
    检测Excel工作表的数据锚点，确定名称列、数据列和数据起始行位置。

    通过查找“读数1”定位读数标题行，计算数据起始行；通过“构件名称”关键词调整名称列，通过“读数1”调整数据列。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
    Returns:
        dict: 锚点信息字典，包含以下键：
            - name_col: 名称列索引（int）
            - data_col: 数据列起始索引（int）
            - data_row: 数据起始行索引（int）
            - read_row: 读数标题行索引（int）
    """
    read_row = None
    for r in range(1, 60):
        for c in range(1, 40):
            if "读数1" in str(ws.cell(row=r, column=c).value or ""):
                read_row = r;
                break
        if read_row: break
    data_start_row = (read_row + 1) if read_row else 7
    name_col = 2
    for r in range(1, (read_row or 15) + 1):
        for c in range(1, 30):
            if "构件名称" in str(ws.cell(row=r, column=c).value or ""):
                name_col = c;
                break
        if name_col != 2: break
    data_col = None
    if read_row:
        for c in range(1, 40):
            if "读数1" in str(ws.cell(row=read_row, column=c).value or ""):
                data_col = c;
                break
    data_col = data_col or 5
    return {"name_col": name_col, "data_col": data_col, "data_row": data_start_row, "read_row": read_row or 6}


def keep_align(cell, value):
    """
    向Excel单元格写入值并保留原有对齐格式，避免格式错乱。

    读取单元格原有对齐方式（水平/垂直对齐、自动换行等），写入值后重新应用这些格式。

    Args:
        cell: Excel单元格对象（openpyxl.cell.cell.Cell）
        value: 待写入的值（str）
    """
    old = cell.alignment or Alignment()
    cell.value = value
    cell.alignment = Alignment(
        horizontal=old.horizontal,
        vertical=old.vertical,
        wrap_text=old.wrap_text,
        textRotation=old.textRotation,
        indent=old.indent,
        shrinkToFit=old.shrinkToFit
    )


def write_block(ws, anchors, pos, item):
    """
    将数据块写入Excel工作表的指定位置，保留格式对齐。

    根据锚点信息计算起始行，写入构件名称和5行数据，确保与模板格式一致。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        anchors: 锚点信息字典（detect_anchors返回结果）
        pos: 数据块在工作表中的位置（0-4，int）
        item: 数据块对象（expand_blocks返回的单个元素）
    """
    r0 = anchors["data_row"] + pos * PER_LINE_PER_BLOCK
    name_col = anchors["name_col"];
    data_col = anchors["data_col"]
    keep_align(ws.cell(row=r0, column=name_col), item["name"])
    for dr in range(PER_LINE_PER_BLOCK):
        for dc in range(9):
            ws.cell(row=r0 + dr, column=data_col + dc).value = item["data"][dr][dc]


def slash_block(ws, anchors, pos):
    """
    用“/”填充Excel工作表中指定位置的数据块，用于补齐未填满的区域。

    在指定位置写入“/”占位符，保留单元格原有对齐格式，确保表格格式统一。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        anchors: 锚点信息字典（detect_anchors返回结果）
        pos: 数据块位置（0-4，int）
    """
    r0 = anchors["data_row"] + pos * PER_LINE_PER_BLOCK
    name_col = anchors["name_col"];
    data_col = anchors["data_col"]
    keep_align(ws.cell(row=r0, column=name_col), "/")
    for dr in range(PER_LINE_PER_BLOCK):
        for dc in range(9):
            ws.cell(row=r0 + dr, column=data_col + dc).value = "/"


def slash_tail(ws, anchors, used_pos):
    """
    用“/”填充工作表中未使用的数据块位置，从已用位置到最后。

    确保工作表数据区域格式统一，未使用的位置明确标记为“/”。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        anchors: 锚点信息字典（detect_anchors返回结果）
        used_pos: 已使用的数据块位置索引（int）
    """
    for rem in range(used_pos, BLOCKS_PER_SHEET):
        slash_block(ws, anchors, rem)



# ===== 元信息固定坐标 =====
def top_left_of_merged(ws, r, c):
    """
    查找合并单元格的左上角单元格坐标，确保值写入正确位置。

    遍历工作表中的合并区域，返回指定单元格所属合并区域的左上角行号和列号。

    Args:
        ws: Excel工作表对象（openpyxl.worksheet.worksheet.Worksheet）
        r: 行号（int）
        c: 列号（int）
    Returns:
        tuple: 左上角单元格的行号和列号（int, int）
    """
    for rng in ws.merged_cells.ranges:
        if rng.min_row <= r <= rng.max_row and rng.min_col <= c <= rng.max_col:
            return rng.min_row, rng.min_col
    return r, c

# ===== 非交互：单日模式导出（供 UI 直接调用） =====
from pathlib import Path
from typing import Union
from openpyxl import load_workbook

def export_single_day_noninteractive(
    src: Union[str, Path],
    meta: dict | None = None,
    single_date: str | None = None,
    *,
    support_strategy: str = "number",
    net_strategy: str = "number",
) -> dict:
    """
    非交互导出（锁定 Mode 3 / 单日）。
    返回: {"excel": Path, "word": Path|None}
    """
    # 0) 校验
    src = Path(str(src)).resolve()
    if not src.exists():
        raise FileNotFoundError(f"未找到 Word 源文件：{src}")

    # 1) 解析 Word
    grouped, categories_present = prepare_from_word(src)

    # 2) 选择 Excel 模板（有支撑版）
    template_path = None
    for name in ("XLSX_WITH_SUPPORT_DEFAULT", "XLSX_TEMPLATE_WITH_SUPPORT", "DEFAULT_XLSX_WITH_SUPPORT"):
        if name in globals() and globals()[name]:
            template_path = Path(globals()[name])
            break
    if not template_path or not Path(template_path).exists():
        raise FileNotFoundError("未找到 Excel 模板常量（XLSX_WITH_SUPPORT_DEFAULT / XLSX_TEMPLATE_WITH_SUPPORT / DEFAULT_XLSX_WITH_SUPPORT）。")

    wb = load_workbook(str(template_path))

    # 3) 执行模式：锁定 Mode 3
    #    —— 注入一次性“非交互日期”，供 run_mode 读取并跳过 ask()
    prev_flag = globals().get("NONINTERACTIVE_MODE3_DATE", None)
    globals()["NONINTERACTIVE_MODE3_DATE"] = single_date if single_date is not None else ""
    try:
        used_pages = run_mode("3", wb, grouped, categories_present)
    finally:
        # run_mode 内部会 pop，这里再兜底清掉
        globals().pop("NONINTERACTIVE_MODE3_DATE", None)
        if prev_flag is not None:
            globals()["NONINTERACTIVE_MODE3_DATE"] = prev_flag

    # 4) 固定元信息、字体与清理
    meta = meta or {}
    apply_meta_fixed(wb, categories_present, meta)
    enforce_mu_font(wb)
    cleanup_unused_sheets(wb, categories_present)

    # 5) 保存到同目录，避免覆盖
    def _unique_name(p: Path) -> Path:
        if not p.exists():
            return p
        stem, suf = p.stem, p.suffix
        i = 1
        while True:
            cand = p.with_name(f"{stem}({i}){suf}")
            if not cand.exists():
                return cand
            i += 1

    out_xlsx = _unique_name(src.parent / "汇总原始记录.xlsx")
    wb.save(str(out_xlsx))

    # 6) 可选：生成 Word 汇总（安全调用）
    word_out = None
    maybe_func = globals().get("export_word_summary", None)
    if callable(maybe_func):
        try:
            word_out = maybe_func(src, grouped)
        except Exception:
            word_out = None

    return {"excel": out_xlsx, "word": word_out}


def apply_meta_fixed(wb, categories_present, meta: dict):
    """
    向Excel工作表写入固定元信息（工程名称、委托编号）到指定位置。

    仅处理目标类型工作表，将工程名称写入C3、委托编号写入L3，支持合并单元格。

    Args:
        wb: Excel工作簿对象（openpyxl.workbook.Workbook）
        categories_present: 存在的构件类型列表（list[str]）
        meta: 元信息字典，含"proj"（工程名称）和"order"（委托编号）键
    """
    for ws in wb.worksheets:
        if not any(ws.title.startswith(p) for p in categories_present): continue

        def _set_rc(r, c, v):
            if not v: return
            r0, c0 = top_left_of_merged(ws, r, c)
            ws.cell(row=r0, column=c0).value = v

        _set_rc(3, 3, meta.get("proj"))  # C3
        _set_rc(3, 12, meta.get("order"))  # L3


def apply_meta_on_pages(wb, pages: list[str], date_str: str):
    """
    向指定 Excel 工作表写入日期元信息，支持公式检测和清除。

    写入逻辑：
    1. 定位到第32行第1列（或其合并单元格的左上角）
    2. 检测该单元格是否包含公式引用
    3. 如果包含公式，先清除公式再写入值（避免多个sheet共享同一数据源）
    4. 写入日期值并保留单元格对齐格式
    5. 输出调试日志以便追踪写入过程

    Args:
        wb: Excel 工作簿对象（openpyxl.workbook.Workbook）
        pages: 工作表名称列表（list[str]）
        date_str: 日期字符串（str），如"2025年1月1日"
    """
    if not pages:
        return
    value = (date_str or "").strip()
    # 调试日志：显示即将写入的日期和页面列表
    if value:
        print(f"\n📅 [apply_meta_on_pages] 写入日期: '{value}' 到 {len(pages)} 个sheet")
    for name in pages:
        if name not in wb.sheetnames:
            continue
        ws = wb[name]
        r0, c0 = top_left_of_merged(ws, 32, 1)
        cell = ws.cell(row=r0, column=c0)

        # 读取当前单元格的值和类型
        old_value = cell.value
        has_formula = False

        # 检测是否包含公式（Excel公式以"="开头）
        if old_value and isinstance(old_value, str) and old_value.startswith('='):
            has_formula = True
            print(f"⚠️  [{name}] 单元格({r0},{c0})包含公式: {old_value}")
            # 先清空单元格，断开公式引用
            cell.value = None

        # 写入日期值（保留对齐格式）
        keep_align(cell, value)

        # 验证写入结果
        actual_value = ws.cell(row=r0, column=c0).value
        if has_formula:
            print(f"✓  [{name}] 已清除公式并写入: '{actual_value}' 到 ({r0},{c0})")
        elif value:
            # 仅在有值时输出日志
            print(f"   [{name}] 写入: '{actual_value}' 到 ({r0},{c0})")


# ===== 规范化 =====
def normalize_date(text: str) -> str:
    """
    将用户输入的环境温度字符串规范化为“X℃”或“X.X℃”格式。

    从输入中提取数字部分（忽略“℃”“度”等符号），整数温度去小数点，小数温度保留有效数字。
    若无法提取有效数字，则返回原始字符串。

    Args:
        text: 用户输入的环境温度字符串（如“24”“24℃”“24.5度”）
    Returns:
        str: 标准化的温度字符串（如“24℃”“24.5℃”）
    """
    s = (text or "").strip()
    if not s: return ""
    if re.fullmatch(r"\d{8}", s):
        y, m, d = int(s[:4]), int(s[4:6]), int(s[6:8]);
        return f"{y}年{m}月{d}日"
    s2 = s.replace("年", " ").replace("月", " ").replace("日", " ")
    for ch in ".-/，,": s2 = s2.replace(ch, " ")
    nums = re.findall(r"\d+", s2)
    if len(nums) >= 3:
        y, m, d = map(int, nums[:3]);
        return f"{y}年{m}月{d}日"
    return s



def _normalize_date_token(tok: str, base_year: int) -> str:
    """将单个日期 token 规范为"YYYY-MM-DD"，失败返回空串。"""
    if not tok:
        return ""
    tok = tok.strip()
    tok = tok.replace("年", "-").replace("月", "-").replace("日", "")
    tok = tok.replace("/", "-").replace(".", "-")
    tok = re.sub(r"\s+", "-", tok)
    if re.fullmatch(r"\d{8}", tok):
        y = int(tok[:4]);
        mth = int(tok[4:6]);
        d = int(tok[6:])
    else:
        m = re.fullmatch(r"(\d{4})-(\d{1,2})-(\d{1,2})", tok)
        if m:
            y, mth, d = map(int, m.groups())
        else:
            m = re.fullmatch(r"(\d{1,2})-(\d{1,2})", tok)
            if not m:
                return ""
            y = base_year
            mth, d = map(int, m.groups())
    if not (1 <= mth <= 12 and 1 <= d <= 31):
        return ""
    return f"{y:04d}-{mth:02d}-{d:02d}"


def _parse_dates_simple(input_str: str):
    """简单解析多个日期，返回 (日期列表, 无效token列表)。"""
    # 允许空格/英文逗号/中文逗号/中文顿号作为分隔
    tokens = [t for t in re.split(r"[,\s，、]+", input_str.strip()) if t]

    res, ignored = [], []
    seen = set()
    base_year = None
    cur_year = datetime.now().year

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        consumed = 1

        # 先尝试把当前 token 当成一个完整日期（支持 8/27、8-27、2025-8-27、2025年8月27日 等等）
        norm = _normalize_date_token(tok, base_year or cur_year)

        if not norm:
            # 尝试 Y M D 这种被空格/逗号拆开的情况：2025 8 27
            if re.fullmatch(r"\d{4}", tok) and i + 2 < len(tokens) \
                    and tokens[i + 1].isdigit() and tokens[i + 2].isdigit():
                norm = _normalize_date_token(
                    f"{tok}-{tokens[i + 1]}-{tokens[i + 2]}",
                    base_year or cur_year
                )
                consumed = 3

            # 尝试 M D：8 27（基于 base_year 或当前年）
            elif tok.isdigit() and i + 1 < len(tokens) and tokens[i + 1].isdigit():
                norm = _normalize_date_token(
                    f"{tok}-{tokens[i + 1]}",
                    base_year or cur_year
                )
                consumed = 2

        if norm:
            # 锁定 base_year，后续 M-D 走同一年
            if base_year is None:
                base_year = int(norm[:4])
            # 去重：同一天不重复计入
            if norm not in seen:
                res.append(norm)
                seen.add(norm)
        else:
            # 记录无法解析的原始 token（或组合）
            ignored.extend(tokens[i:i + consumed])

        i += consumed

    return res, ignored

    # ===== 交互 =====
HELP_HOME = f"""
==================== 原始记录自动填写程序 | 帮助中心（{VERSION}） ====================

一、基本流程
  1) 在“请输入 Word 源路径”处输入 .docx 文件路径（输入 help 打开本帮助）；
  2) 程序读取源文件并生成《汇总原始记录.docx》；
  3) 选择模式 1 / 2 / 3 / 4，按向导完成分配与出表。

二、全局操作
  • 任意步骤输入小写 q 返回上一步；仅在路径输入界面输入大写 Q 退出程序。
  • 在路径输入界面输入 help 打开本帮助；回车返回路径输入。

三、输入规范（程序自动标准化）
  • 日期：
      支持以下任意形式，自动规范：YYYY-MM-DD / YYYY/MM/DD / YYYY.MM.DD /
      YYYY MM DD / YYYYMMDD / M-D / M/D / M.D / M D / YYYY年M月D日。
  • 特殊指令：
      *   表示“全部接收”（如网架子类编号范围）；
      lk  表示“留空，不接收”（仅网架范围录入时可用）；
      a   在“分配确认”阶段表示“将未分配构件并入最后一天”。

四、分类与模板、排序规则
  • 已内置类别：钢柱、钢梁、支撑（WZ）、网架（含 XX/FG/SX/泛称）；未识别归“其他”（复用钢柱模板）。
  • μ 值识别：任一读数出现“≥4 位连续数字”或“绝对值≥1000 的纯数值”即判定为 μ。
      - 同一桶内：先写普通页，后写 μ 页；同页不混。
      - 跨桶编号连续（普通+μ 统一流水号），不为 μ 额外补建“空普通页”。
      - 未使用的“裸 μ 模板页”（如“钢梁μ”）会在出表后自动清理。
  • 页池命名：沿用模板页名，不将日期/楼层拼入 Sheet 名称。
  • 楼层排序：地下 B* → 数字层（1F↑）→ 机房层 → 屋面；同层内按“WZ 编号 → 名称中的数字 → 字典序”。

五、支撑/网架分桶策略（Mode 1/2/3 将先询问）
  1 = 按编号（WZ 号/网架子类编号）； 2 = 按楼层（与钢柱/钢梁一致）。
  注：网架支持为各子类单独配置范围；同一日期写入同一张“网架”表。

六、使用提示
  • 无论成功或失败，流程结束都会回到路径输入；仅在路径输入处输入大写 Q 才会退出。
  • 运行前务必关闭相关 Word/Excel 文件，避免“文件占用”导致读写失败。

—— 常见问题快速排查 ——
1) “Excel 被占用/无法保存”：关闭模板或目标文件后重试。
2) “找不到文件”：检查路径是否含引号/空格；或使用默认路径。
3) “未识别到数据表”：源 Word 表格需包含“测点1”“平均值”表头。
4) “缺少 μ 模板”：确保模板包含所需 μ 页（钢柱μ/钢梁μ/支撑μ/网架μ）。
5) “网架全部进入 μ 页”：核查是否确有 4 位数或 ≥1000 的读数触发阈值。

如需查看各模式说明，请在下方输入 1 / 2 / 3 / 4；回车或 q 返回。

=====================================================================
"""

HELP_TEXTS = {
    "1": r"""====================  Mode 1 | 按日期分桶（默认稳健）  ====================

适用场景
  将全部构件分配至多个日期；支持“后面的日期优先”（默认）或“前面的日期优先”。

操作流程
  1) 选择模式：输入 1；
  2) 存在“支撑/网架”时，先选择分桶策略：1=按编号，2=按楼层；
  3) （仅网架）为各子类录入编号范围：回车=沿用上次；*=全部；lk=留空不接收；
  4) 输入 1–10 个日期，程序自动去重并规范化格式；
  5) 冲突处理：回车=“后面的日期优先”（默认），n=“前面的日期优先”；
  6) 预览分配：回车=确认生成，n=取消，a=将未分配并入最后一天；
  7) 执行出表：批量写入页池及元信息（日期）。

要点说明
  • μ 判定依据“≥4 位数字”或“绝对值≥1000”的读数；同桶内“普通在前、μ 在后”，跨桶流水号连续。
  • 支撑/网架按所选策略参与排序，和钢柱/钢梁并行处理。

返回/退出：任意步骤输入 q 返回上一步；仅路径输入界面输入大写 Q 退出。
=====================================================================
""",

    "2": r"""====================  Mode 2 | 按楼层断点（按层出报）  ====================

适用场景
  通过断点将楼层划分为多个“楼层桶”（如 1F–3F、4F–6F、B3–B1、机房层、屋面），
  每个楼层桶对应一个日期（可相同）

操作流程
  1) 选择模式：输入 2；
  2) 如存在“支撑/网架”，先选择分桶策略（1=编号 / 2=楼层）；
  3) 录入断点（如“5 10”），排序规则固定为“B* → 1F↑ → 机房层 → 屋面”；
  4) 为每个桶设置日期；
  5) 预览并确认，执行出表。

要点说明
  • μ/普通分流及编号规则与全局一致；Sheet 命名不包含日期/楼层信息。

返回/退出：任意步骤输入 q 返回上一步；仅路径输入界面输入大写 Q 退出。
=====================================================================
""",

    "3": r"""====================  Mode 3 | 单日模式（快速制表）  ====================

适用场景
  全量数据同一日期出报，或需要快速生成成表。

操作流程
  1) 选择模式：输入 3；
  2) 如存在“支撑/网架”，先选择分桶策略（1=编号 / 2=楼层）；
  3) 程序按“每页 5 组 × 每组 5 行”自动分页，写入页池与元信息。

要点说明
  • μ 判定与编号规则同全局；同桶“普通先、μ后”，流水号跨桶连续。
  • 支撑/网架遵循所选策略并入整体排序。

返回/退出：任意步骤输入 q 返回上一步；仅路径输入界面输入大写 Q 退出。
=====================================================================
""",

    "4": r"""================  Mode 4 | 楼层 × 日期 切片（均分 / 配额）  ================

适用场景
  同一楼层需要分配到多天；支持“均分”或“每日上限（配额）”。

核心概念
  • 共用计划：为一批楼层配置同一组日期与每日上限（留空上限=均分）；
  • 默认计划（*）：对“未单独配置”的楼层生效的兜底方案；
  • 若仍有剩余：可统一指定日期/温度，或返回 Mode 1 继续分桶。

操作流程
  1) 选择模式：输入 4；
  2) 选择本次涉及的楼层（回车=全部识别到的楼层；支持 B2/5F/机房层/屋面等标签）；
  3) 是否采用“共用计划”：y=对所选楼层共用；回车=逐楼层设置；
  4) 如仍有未配置楼层：可创建默认计划（*）；
  5) 执行切片、分页与写入；必要时选择兜底策略。

要点说明
  • 楼层排序及 μ/普通分流、编号规则继承全局逻辑。
  • 网架/支撑继续按所选策略参与切片和出表。

返回/退出：任意步骤输入 q 返回上一步；仅路径输入界面输入大写 Q 退出。
=====================================================================
""",
}





def tutorial_browser():
    """显示模式教程浏览器。"""
    print(HELP_HOME)
    viewed = False
    while True:
        prompt = "还要查看其他模式？输入 1/2/3/4，回车或 q 返回。\n→ " if viewed else "查看哪个模式？输入 1/2/3/4，回车或 q 返回路径输入。\n→ "
        sel = input(prompt).strip()
        if sel in ("", "q"):
            return
        if sel in HELP_TEXTS:
            print(HELP_TEXTS[sel])
            viewed = True
        else:
            print("仅接受 1/2/3/4 或回车/q。")


def prompt_path(prompt, default: Path) -> Path:
    """
    交互式获取用户输入文件路径，验证文件存在性并返回有效路径。

    提示用户输入文件路径，支持直接回车使用默认路径；自动处理路径中的引号；
    若输入路径无效（文件不存在），则显示错误提示并重新请求输入，确保返回有效文件路径。

    Args:
        prompt: 路径输入提示信息（str）
        default: 默认文件路径（Path对象）
    Returns:
        Path: 经过验证的有效文件路径
    """
    while True:
        raw = ask(f"{prompt}（回车默认：{default}）")
        if raw.lower() == "help":
            tutorial_browser()
            continue
        p = Path(raw.strip('"')) if raw else default
        if p.exists() and p.is_file():
            return p
        print(f"❌ 找不到文件：{p}")


def prompt_floor_breaks(label: str):
    """
    交互式获取楼层断点列表，支持无效输入并返回空值处理。

    提示用户输入空格分隔的楼层断点（如"5 10"），支持直接回车表示不分段；
    自动过滤重复值并按升序排序；若输入无效（非数字）则返回空列表。

    Args:
        label: 提示信息前缀（str）
    Returns:
        list[int]: 排序后的楼层断点列表（空列表表示不分段）
    """
    txt = ask(f"{label} 断点楼层（空格分隔，如 5 10；回车=不分段）：")
    if not txt: return []
    try:
        return sorted({int(x) for x in txt.split()})
    except:
        return []


# ===== 日期分桶（泛化到任意类别） =====
import re
from typing import List, Tuple

def _parse_int_ranges(expr: str) -> List[Tuple[int, int]]:
    """
    将楼层选择表达式解析为一组闭区间 [(lo, hi), ...]。

    约定（与需求对齐）：
      - 留空（空字符串/只空白） => **不要任何楼层**（返回一个永不命中的“空集哨兵”区间）
      - '*'（或全角'＊'）      => **全都要**（返回 []，让上游的空列表=全包含逻辑生效）

    支持：
      - 数字：'3'
      - 范围：'1-10'（连字符可为 - — – ~ 至 到）
      - 特殊楼层：机房/机房层/JF/jf，屋面/屋顶层/顶层/WM/wm/roof
      - 混合：'10-机房'、'机房-屋面'、'8-屋面'
      - 多分隔：空格/英文逗号/中文逗号/顿号/分号等
    """
    text = (expr or "").strip()

    # —— 规则：留空 = 不要任何楼层（返回一个永不命中的区间，避免上游把空当“全包含”）——
    if not text:
        return [(1, 0)]  # lo > hi，任何楼层都不会命中；且列表非空，不会触发“全包含”

    # '*' = 全都要：保持空列表，让上游的“空=全包含”生效
    if text in ("*", "＊"):
        return []

    # 先做 NFKC 规范化 + 统一连字符
    text = unicodedata.normalize("NFKC", text)
    # 把各种“看起来像连字符/波浪线/中文至到”统一成 '-'
    text = re.sub(r"[－—–−~～〜至到]", "-", text)

    # ***** 特殊楼层映射 *****
    JF_VAL = 10**6 - 1    # 机房
    WM_VAL = 10**6        # 屋面/屋顶层/顶层
    SPECIAL_MAP = {
        "机房": JF_VAL, "机房层": JF_VAL, "jf": JF_VAL,
        "屋面": WM_VAL, "屋顶层": WM_VAL, "顶层": WM_VAL,
        "wm": WM_VAL, "roof": WM_VAL,
    }

    def norm_token(tok: str) -> str:
        return tok.strip().lower()

    # 分词：空格、英文/中文逗号、顿号、分号
    tokens = [t for t in re.split(r"[,\uFF0C\u3001;\uFF1B\s]+", text) if t.strip()]

    ranges: List[Tuple[int, int]] = []

    # 各类正则
    re_int = re.compile(r"^\s*\d+\s*$")
    # 统一后只需要匹配 '-'
    re_num_num = re.compile(r"^\s*(\d+)\s*-\s*(\d+)\s*$")
    re_a_sp   = re.compile(r"^\s*(\d+)\s*-\s*([^\d\s]+)\s*$")
    re_sp_b   = re.compile(r"^\s*([^\d\s]+)\s*-\s*(\d+)\s*$")
    re_sp_sp  = re.compile(r"^\s*([^\d\s]+)\s*-\s*([^\d\s]+)\s*$")

    def sp_val(s: str):
        key = norm_token(s).replace("（", "(").replace("）", ")")
        return SPECIAL_MAP.get(key)

    for raw in tokens:
        tok = raw.strip()
        if not tok:
            continue

        # 单个数字
        if re_int.match(tok):
            v = int(tok)
            ranges.append((v, v))
            continue

        # 数字-数字
        m = re_num_num.match(tok)
        if m:
            a, b = int(m.group(1)), int(m.group(2))
            if a > b:
                a, b = b, a
            ranges.append((a, b))
            continue

        # 数字-特殊
        m = re_a_sp.match(tok)
        if m:
            a = int(m.group(1))
            rb = sp_val(m.group(2))
            if rb is not None:
                lo, hi = (a, rb) if a <= rb else (rb, a)
                ranges.append((lo, hi))
                continue

        # 特殊-数字
        m = re_sp_b.match(tok)
        if m:
            la = sp_val(m.group(1))
            b = int(m.group(2))
            if la is not None:
                lo, hi = (la, b) if la <= b else (b, la)
                ranges.append((lo, hi))
                continue

        # 特殊-特殊
        m = re_sp_sp.match(tok)
        if m:
            la, lb = sp_val(m.group(1)), sp_val(m.group(2))
            if la is not None and lb is not None:
                lo, hi = (la, lb) if la <= lb else (lb, la)
                ranges.append((lo, hi))
                continue

        # 单个特殊词
        sv = sp_val(tok)
        if sv is not None:
            ranges.append((sv, sv))
            continue

        # 未识别片段，友好提示（不影响运行）
        print(f"[hint] 未识别片段：{raw}（已忽略，不参与楼层筛选）")

    # 如果仍然什么都没解析到：给一个“空集哨兵”，避免被误认为“全包含”
    if not ranges:
        return [(1, 0)]

    # 合并区间
    ranges.sort(key=lambda x: (x[0], x[1]))
    merged: List[Tuple[int, int]] = []
    for lo, hi in ranges:
        if not merged:
            merged.append((lo, hi))
        else:
            mlo, mhi = merged[-1]
            if lo <= mhi:         # 重叠/相邻都并
                merged[-1] = (mlo, max(mhi, hi))
            else:
                merged.append((lo, hi))
    return merged



def parse_rule(text: str):
    """
    解析数据分发规则字符串为结构化规则字典。

    支持两种规则类型：
    - 启用所有数据：输入“*”“all”“全部”“所有”时，返回启用状态且空范围（表示接收所有数据）
    - 范围规则：其他输入解析为整数范围列表（通过_parse_int_ranges处理）

    Args:
        text: 规则字符串（如“*”“1-3 5”“全部”）
    Returns:
        dict: 规则字典，包含：
            - enabled: 是否启用该规则（bool）
            - ranges: 解析后的范围列表（list[tuple[int, int]]，空列表表示全部）
    """
    s = (text or "").strip()
    if not s:
        return {"enabled": False, "ranges": []}
    if _is_explicit_all_token(s):
        return {"enabled": True, "ranges": [], "explicit_all": True}
    return {"enabled": True, "ranges": _parse_int_ranges(s)}


def _is_lk(s: str) -> bool:
    """大小写及全角半角均识别 'lk'。"""
    return unicodedata.normalize('NFKC', (s or '')).strip().lower() == 'lk'


_STAR_TOKENS = {"*", "全部", "所有"}


def _is_explicit_all_token(value) -> bool:
    """判定输入是否表示显式的“全部接收”。"""
    if not isinstance(value, str):
        return False
    token = unicodedata.normalize("NFKC", value or "").strip()
    if not token:
        return False
    if token in _STAR_TOKENS:
        return True
    return token.casefold() == "all"


def _in_ranges(val: int, ranges):
    """
    判断值是否在指定的范围列表内，支持空范围表示“全部包含”。

    范围列表为空时默认包含所有值；否则检查值是否落在任一范围的闭区间内。

    Args:
        val: 待判断的整数（如楼层号、支撑编号）
        ranges: 范围元组列表（如[(1,3), (5,7)]），空列表表示全部
    Returns:
        bool: 在范围内返回True，否则返回False
    """
    if ranges is None: return False
    if ranges == []: return True  # noqa
    for a, b in ranges:
        if a <= val <= b: return True
    return False


def net_part(name: str) -> str:
    """
    返回 'XX' / 'FG' / 'SX' / 'GEN'（泛称）之一；大小写不敏感，兼容中文别名。
    """
    s = name.upper()
    # 形如 XX1 / XX-12 / XX_003，或中文别名
    if re.search(r"(?<![A-Z0-9])XX(?=[-_]?\d+)|下\s*弦", s):
        return "XX"
    if re.search(r"(?<![A-Z0-9])FG(?=[-_]?\d+)|腹\s*杆", s):
        return "FG"
    if re.search(r"(?<![A-Z0-9])SX(?=[-_]?\d+)|上\s*弦", s):
        return "SX"
    if re.search(r"\bWJ\b|网\s*架|SPACE\s*FRAME|GRID", s):
        return "GEN"
    return "GEN"


def _net_no(name: str):
    """
    从网架构件名里提取编号（XX12 / FG-03 / SX_7 / 网架-15 等）。

    仅在明确前缀或泛称存在时才解析编号，避免误吃其他数字。
    """
    s = name.upper()
    part = net_part(name)
    if part in ("XX", "FG", "SX"):
        m = re.search(rf"{part}\s*[-_]?(\d+)", s)
        return int(m.group(1)) if m else None
    m = re.search(r"(?:WJ|网架|SPACE\s*FRAME|GRID)\s*[-_]?(\d+)", s)
    return int(m.group(1)) if m else None


def _wz_no(name: str):
    """
    从支撑构件名称中提取编号（如从“WZ3”“支撑-5”中提取3、5）。

支持关键词匹配：
- 含“WZ”或“ZC”前缀（如“WZ12”“ZC-8”）
- 含“支撑”关键词（如“支撑6”“斜撑-3”）
提取失败时返回None。

Args:
    name: 支撑构件名称字符串（如“WZ5”“支撑-10”）
Returns:
    int | None: 提取的编号，失败则返回None
"""
    m = re.search(r"(?i)\b(?:WZ|ZC)\s*[-–—]?\s*(\d+)\b", name)
    if m: return int(m.group(1))
    m = re.search(r"支撑\s*[-–—]?\s*(\d+)", name)
    return int(m.group(1)) if m else None


def _match_keywords(name: str, kws):
    """
    判断构件名称是否包含任意关键词（忽略大小写）。

    关键词列表为空时默认匹配所有名称；否则检查名称是否含任一关键词（不区分大小写）。

    Args:
        name: 构件名称字符串
        kws: 关键词列表（如["3层", "东立面"]）
    Returns:
        bool: 包含任一关键词返回True，否则返回False（关键词为空时返回True）
    """
    if not kws: return True
    s = name.lower()
    return any(k.lower() in s for k in kws)


def prompt_mode():
    """模式选择，支持 q 返回。"""
    txt = ask("模式选择：1) 按日期分桶  2) 按楼层断点  3) 单日模式  4) 楼层+日期配额")
    if txt in ("", "1"):
        return "1"
    if txt in ("2", "3", "4"):
        return txt
    return "1"


def prompt_bucket_priority():
    """询问规则重叠优先级。"""
    ans = ask("规则重叠将按【后面的天】优先并自动做差（回车=是 / n=否）：", lower=True)
    return ans != 'n'


def prompt_later_priority():
    """供适配层覆盖的“后桶优先”询问接口。"""
    return prompt_bucket_priority()


def prompt_auto_merge_remains(*_, **__):
    """供适配层覆盖的“是否自动并入剩余构件”接口。默认返回 ``None``。"""
    return None


def prompt_keywords_for_bucket(*_, **__):
    """供适配层覆盖的关键词输入接口。默认返回 ``None`` 表示继续交互询问。"""
    return None


def prompt_support_strategy_for_bucket():
    """在需要支撑分桶策略时询问一次。"""
    global support_bucket_strategy
    if support_bucket_strategy is None:
        ans = ask("支撑分桶方式：1) 按编号 2) 按楼层（回车=1）")
        support_bucket_strategy = "floor" if ans == "2" else "number"
    return support_bucket_strategy


def prompt_net_strategy_for_bucket():
    """在需要网架分桶策略时询问一次。"""
    global net_bucket_strategy
    if net_bucket_strategy is None:
        ans = ask("网架分桶方式：1) 按编号  2) 按楼层（回车=1）")
        net_bucket_strategy = "floor" if ans == "2" else "number"
    return net_bucket_strategy


def detect_net_parts_for_category(grouped, cat="网架"):
    """检测本次运行实际出现的网架子类集合。"""
    parts = set()
    for g in grouped.get(cat, []):
        parts.add(net_part(g["name"]))
    return parts or {"GEN"}


def prompt_date_buckets(categories_present, grouped):
    """
    交互式收集日期桶配置，支持1-10天的检测数据分发规则。

    为每天配置：
    - 日期（自动标准化为“YYYY年MM月DD日”）
    - 环境温度（自动标准化为“X℃”）
    - 各构件类型的接收规则（楼层/编号范围）
    - 关键词筛选（可选）

    Args:
        categories_present: 存在的构件类型列表（如["钢柱", "支撑"]）
        grouped: 按类型分组的构件数据，用于检测网架子类
    Returns:
        list[dict]: 日期桶配置列表，每个元素含日期、环境、规则等信息
    """
    while True:
        n_txt = ask("共有几天（1-10，回车=1）：")
        if not n_txt: n = 1; break
        if n_txt.isdigit() and 1 <= int(n_txt) <= 10:
            n = int(n_txt);
            break
        print("请输入 1-10 之间的整数。")
    buckets = []

    # 预先检测网架子类，向用户提示出现的名称
    net_parts_present = set()
    if "网架" in categories_present:
        net_parts_present = detect_net_parts_for_category(grouped, "网架")
        if net_parts_present:
            name_map = {"XX": "XX", "FG": "FG", "SX": "SX", "GEN": "泛称"}
            pretty = "、".join(name_map.get(p, p) for p in sorted(net_parts_present))
            print(f"🕸 本次识别到的网架名称：{pretty}")

    for i in range(1, n + 1):
        print(f"\n—— 第 {i} 天 ——")
        d = ask("📅 日期（20250101 / 2025年1月1日 / 2025 1 1 / 2025.1.1 / 2025-1-1 / 1-1 / 01-01）：")

        rules = {}
        for cat in categories_present:
            if cat == "支撑":
                prompt_support_strategy_for_bucket()
                if support_bucket_strategy == "floor":
                    txt = ask("🦾 支撑 楼层规则（例：1-3 5 7-10 屋面；留空=不接收；*=不限）：")
                else:
                    txt = ask("🦾 支撑 编号范围（例：1-12 20-25；留空=不接收；*=不限）：")
                rules[cat] = parse_rule(txt)
            elif cat == "网架":
                prompt_net_strategy_for_bucket()
                present_parts = net_parts_present
                sub_rules = {}
                if net_bucket_strategy == "number":
                    prev_rule = None
                    for part in sorted(present_parts - {"GEN"}):
                        placeholder = "同上" if prev_rule else "不接收"
                        txt = ask(f"🕸 网架-{part} 编号范围（例：1-12 20-25；留空={placeholder}；*=所有；lk=不接收）：")
                        if _is_lk(txt):
                            sub_rules[part] = {"enabled": False, "ranges": []}
                        elif txt == "":
                            sub_rules[part] = prev_rule or {"enabled": False, "ranges": []}
                        else:
                            sub_rules[part] = parse_rule(txt)
                        print(f"✅ 已设置 网架-{part}: {sub_rules[part]}")
                        prev_rule = sub_rules[part] if txt != "" else prev_rule
                    if "GEN" in present_parts:
                        placeholder = "同上" if prev_rule else "不接收"
                        txt = ask(f"🕸 网架-泛称 编号范围（留空={placeholder}；*=所有；lk=不接收）：")
                        if _is_lk(txt):
                            sub_rules["GEN"] = {"enabled": False, "ranges": []}
                        elif txt == "":
                            sub_rules["GEN"] = prev_rule or {"enabled": False, "ranges": []}
                        else:
                            sub_rules["GEN"] = parse_rule(txt)
                        print(f"✅ 已设置 网架-泛称: {sub_rules['GEN']}")
                        prev_rule = sub_rules["GEN"] if txt != "" else prev_rule
                else:
                    for part in sorted(present_parts - {"GEN"}):
                        txt = ask(f"🕸 网架-{part} 楼层规则（例：1-3 5 7-10 屋面；留空=不接收；*=不限；lk=不接收）：")
                        if _is_lk(txt):
                            rule = {"enabled": False, "ranges": []}
                        else:
                            rule = parse_rule(txt)
                        sub_rules[part] = rule
                        print(f"✅ 已设置 网架-{part}: {rule}")
                    if "GEN" in present_parts:
                        txt = ask("🕸 网架-泛称 楼层规则（留空=不接收；*=不限；lk=不接收）：")
                        if _is_lk(txt):
                            rule = {"enabled": False, "ranges": []}
                        else:
                            rule = parse_rule(txt)
                        sub_rules["GEN"] = rule
                        print(f"✅ 已设置 网架-泛称: {rule}")
                rules[cat] = {"strategy": net_bucket_strategy, "parts": sub_rules}
            else:
                txt = ask(f"🏗 {cat} 楼层规则（例：1-3 5 7-10 屋面；留空=不接收；*=不限）：")
                rules[cat] = parse_rule(txt)
        kws_prefilled = prompt_keywords_for_bucket(
            bucket_index=i - 1,
            rules=rules,
            categories_present=categories_present,
        )
        if kws_prefilled is None:
            kws_txt = ask("🔎 关键词（可多个，空格/逗号分隔；留空=无需）：")
            kws = [k for k in re.split(r"[,\s，]+", kws_txt) if k] if kws_txt else []
        else:
            if isinstance(kws_prefilled, str):
                kws = [k for k in re.split(r"[,\s，]+", kws_prefilled) if k]
            else:
                kws = [str(k).strip() for k in kws_prefilled if str(k).strip()]
        buckets.append({
            "date_raw": d,
            "date": normalize_date(d) if d else "",
            "rules": rules,
            "kws": kws
        })
    return buckets


def assign_by_buckets(cat_groups: dict, buckets, later_priority=True):
    """
    将构件数据组按日期桶规则分配到对应天数，支持规则重叠处理。

    分配逻辑：
    1. 按构件类型遍历数据组
    2. 根据日期桶规则（楼层/编号范围+关键词）匹配数据
    3. 规则重叠时按“后定义桶优先”（可通过参数关闭）
    返回分配结果和未匹配的数据。

    Args:
        cat_groups: 按类型分组的构件数据（键为类型，值为数据组列表）
        buckets: 日期桶配置列表（prompt_date_buckets返回结果）
        later_priority: 规则重叠时是否后定义桶优先，默认True
    Returns:
        tuple: 包含两个元素的元组：
            - cat_byb: 按类型和桶分配的结果（dict[类型][桶索引] = 数据组列表）
            - remain_by_cat: 未分配的数据（dict[类型] = 数据组列表）
    """
    # 输出：cat_byb[cat][bucket_index] = [groups...];  remain_by_cat[cat] = [groups...]
    cat_byb = {cat: {i: [] for i in range(len(buckets))} for cat in cat_groups}
    assigned = {cat: set() for cat in cat_groups}
    order = range(len(buckets) - 1, -1, -1) if later_priority else range(len(buckets))
    sup_strategy = (support_bucket_strategy or "number") if support_bucket_strategy else "number"
    sup_strategy = sup_strategy.lower()
    net_strategy_default = (net_bucket_strategy or "number") if net_bucket_strategy else "number"
    net_strategy_default = net_strategy_default.lower()
    for cat, groups in cat_groups.items():
        for idx, g in enumerate(groups):
            # 计算匹配
            fl = floor_of(g["name"])
            wzno = _wz_no(g["name"]) if cat == "支撑" and sup_strategy == "number" else None
            for bi in order:
                b = buckets[bi]
                bucket_rules = (b or {}).get("rules") or {}
                rule = bucket_rules.get(cat)
                if not rule:
                    continue
                if cat != "网架" and not rule.get("enabled"):
                    continue
                ok = False  # noqa
                if cat == "支撑":
                    if sup_strategy == "number":
                        rng = rule["ranges"]
                        ok_num = True if rng == [] else (wzno is not None and _in_ranges(wzno, rng))
                        ok = ok_num
                    else:
                        ok = _in_ranges(fl, rule["ranges"])
                elif cat == "网架":
                    parts = (rule or {}).get("parts") or {}
                    part = net_part(g["name"])
                    part_rule = parts.get(part) or parts.get("GEN")
                    if not (part_rule and part_rule.get("enabled")):
                        continue
                    bucket_net_strategy = (rule.get("strategy") or net_strategy_default).lower()
                    if bucket_net_strategy == "number":
                        no = _net_no(g["name"])
                        ok = (no is not None) and _in_ranges(no, part_rule["ranges"])
                    else:
                        ok = _in_ranges(fl, part_rule["ranges"])
                else:
                    ok = _in_ranges(fl, rule["ranges"])
                kws_list = b.get("kws") if isinstance(b, dict) else None
                if ok and _match_keywords(g["name"], kws_list):
                    cat_byb[cat][bi].append(g)
                    assigned[cat].add(idx)
                    break

    remain_by_cat = {cat: [g for i, g in enumerate(groups) if i not in assigned[cat]]
                     for cat, groups in cat_groups.items()}
    return cat_byb, remain_by_cat


def _to_bool(x):
    if isinstance(x, bool):
        return x
    s = str(x).strip().lower()
    return s in {"1", "true", "y", "yes", "on"}


class Mode1ConfigProvider:
    """前端配置适配层，提供 Mode 1 所需的结构化配置。"""

    def __init__(
            self,
            buckets,
            support_strategy,
            net_strategy,
            later_priority,
            auto_merge_rest,
            meta=None,
    ):
        self.raw_buckets = list(buckets or [])
        self.support_strategy = (support_strategy or "number").lower()
        self.net_strategy = (net_strategy or "number").lower()
        self.later_priority = _to_bool(later_priority)
        self.auto_merge_rest = _to_bool(auto_merge_rest)
        self.meta = dict(meta or {})
        self._normalized_buckets = [self._normalize_bucket(b) for b in self.raw_buckets]

    def _normalize_bucket(self, bucket):
        data = dict(bucket or {})
        date_raw = data.get("date_raw") or data.get("date") or ""
        kws = self._normalize_keywords(data.get("kws"))
        normalized = {
            "date_raw": date_raw,
            "date": normalize_date(date_raw) if date_raw else "",
            "rules": {},
            "kws": kws,
        }
        rules_in = data.get("rules") or data.get("parts") or {}
        for cat, rule in rules_in.items():
            if cat == "网架":
                normalized["rules"][cat] = self._normalize_net_rule(rule)
            else:
                normalized_rule = self._normalize_simple_rule(rule)
                if normalized_rule:
                    normalized["rules"][cat] = normalized_rule
        return normalized

    def _normalize_keywords(self, kws):
        if not kws:
            return []
        if isinstance(kws, str):
            parts = [k for k in re.split(r"[,\s，]+", kws) if k]
            return parts
        parts = []
        for item in kws:
            s = str(item).strip()
            if s:
                parts.append(s)
        return parts

    def _normalize_simple_rule(self, rule_data):
        if rule_data is None:
            return {"enabled": False, "ranges": []}
        explicit_all = False
        if isinstance(rule_data, dict):
            enabled_flag = rule_data.get("enabled")
            if enabled_flag is None:
                enabled_flag = True
            ranges_raw = rule_data.get("ranges")
            explicit_all = bool(
                rule_data.get("explicit_all")
                or _is_explicit_all_token(rule_data.get("raw"))
                or _is_explicit_all_token(rule_data.get("text"))
                or _is_explicit_all_token(ranges_raw if isinstance(ranges_raw, str) else "")
                )
        else:
            enabled_flag = True
            ranges_raw = rule_data
            explicit_all = _is_explicit_all_token(rule_data if isinstance(rule_data, str) else "")
        ranges = self._coerce_ranges(ranges_raw)
        if ranges is None:
            return {"enabled": False, "ranges": []}
        if ranges == [(1, 0)]:
            return {"enabled": False, "ranges": []}
        if ranges == [] and not explicit_all:
            return {"enabled": False, "ranges": []}
        return {"enabled": bool(enabled_flag), "ranges": ranges}

    def _coerce_ranges(self, ranges_raw):
        if isinstance(ranges_raw, list):
            return list(ranges_raw)
        s = unicodedata.normalize("NFKC", str(ranges_raw or "")).strip()
        if not s:
            return None
        if _is_lk(s):
            # 返回一个“空集哨兵”，上层据此把 enabled 置 False
            return [(1, 0)]
        if _is_explicit_all_token(s):
            return []
        return _parse_int_ranges(s)

    def _normalize_net_rule(self, rule):
        data = dict(rule or {})
        strategy = (data.get("strategy") or self.net_strategy or "number").lower()
        parts_in = data.get("parts") or {}
        parts_out = {}
        for part, part_rule in parts_in.items():
            parts_out[part] = self._normalize_simple_rule(part_rule)
        return {"strategy": strategy, "parts": parts_out}

    def get_buckets(self):
        """返回深拷贝的规范化桶配置，供 run_mode 使用。"""
        return copy.deepcopy(self._normalized_buckets)


def run_mode1_with_provider(src_docx, out_dir, provider: "Mode1ConfigProvider"):
    """以适配层提供的数据运行 Mode1，无需交互。"""

    if provider is None:
        raise ValueError("provider 不能为空")

    src = Path(str(src_docx)).resolve()
    if not src.exists():
        raise FileNotFoundError(f"未找到 Word 源文件：{src}")

    out_dir = Path(out_dir) if out_dir is not None else src.parent
    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    cache_src = _PROBE_CACHE.get("src")
    grouped = None
    categories_present = None
    if cache_src and Path(str(cache_src)).resolve() == src:
        grouped = _PROBE_CACHE.get("grouped") or defaultdict(list)
        categories_present = list(_PROBE_CACHE.get("categories") or [])
    if grouped is None or categories_present is None:
        grouped, categories_present = prepare_from_word(src)
    else:
        if not isinstance(grouped, defaultdict):
            tmp = defaultdict(list)
            for k, v in (grouped or {}).items():
                tmp[k] = list(v)
            grouped = tmp

    tpl_path = XLSX_WITH_SUPPORT_DEFAULT
    if not tpl_path.exists():
        raise FileNotFoundError(f"Excel 模板不存在：{tpl_path}")
    wb = load_workbook_safe(tpl_path)

    buckets = provider.get_buckets()

    _pd = globals().get("prompt_date_buckets")
    _ps = globals().get("prompt_support_strategy_for_bucket")
    _pn = globals().get("prompt_net_strategy_for_bucket")
    _pl = globals().get("prompt_later_priority")
    _pa = globals().get("prompt_auto_merge_remains")
    _pk = globals().get("prompt_keywords_for_bucket")

    def __pd(*_, **__):
        return buckets

    def __ps(*_, **__):
        set_support_strategy(provider.support_strategy)
        return provider.support_strategy

    def __pn(*_, **__):
        set_net_strategy(provider.net_strategy)
        return provider.net_strategy

    def __pl(*_, **__):
        return provider.later_priority

    def __pa(*_, **__):
        return provider.auto_merge_rest

    def __pk(*_, **__):
        return []

    try:
        globals()["prompt_date_buckets"] = __pd
        globals()["prompt_support_strategy_for_bucket"] = __ps
        globals()["prompt_net_strategy_for_bucket"] = __pn
        globals()["prompt_later_priority"] = __pl
        globals()["prompt_auto_merge_remains"] = __pa
        globals()["prompt_keywords_for_bucket"] = __pk

        used_pages = run_mode(
            "1",
            wb,
            categories_present=categories_present,
            grouped_preloaded=grouped,
        )
    finally:
        if _pd is not None:
            globals()["prompt_date_buckets"] = _pd
        if _ps is not None:
            globals()["prompt_support_strategy_for_bucket"] = _ps
        if _pn is not None:
            globals()["prompt_net_strategy_for_bucket"] = _pn
        if _pl is not None:
            globals()["prompt_later_priority"] = _pl
        if _pa is not None:
            globals()["prompt_auto_merge_remains"] = _pa
        if _pk is not None:
            globals()["prompt_keywords_for_bucket"] = _pk
        set_support_strategy(None)
        set_net_strategy(None)

    meta = provider.meta or {}
    apply_meta_fixed(wb, categories_present, meta)
    enforce_mu_font(wb)
    cleanup_unused_sheets(wb, used_pages, bases=tuple(CATEGORY_ORDER))

    def _unique_out_path(dest_dir: Path, stem: str) -> Path:
        cand = dest_dir / f"{stem}.xlsx"
        if not cand.exists():
            return cand
        i = 1
        while True:
            cand = dest_dir / f"{stem}({i}).xlsx"
            if not cand.exists():
                return cand
            i += 1

    final_path = _unique_out_path(out_dir, f"{TITLE}_报告版")
    save_workbook_safe(wb, final_path)

    word_out = src.with_name("汇总原始记录.docx")
    if not word_out.exists():
        all_rows = _PROBE_CACHE.get("all_rows")
        if all_rows:
            try:
                doc_out = build_summary_doc_with_progress(all_rows)
                set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
                save_docx_safe(doc_out, word_out)
            except Exception:
                pass
    return final_path, word_out


def set_support_strategy(strategy: str | None):
    """设置全局支撑分桶策略。"""
    global support_bucket_strategy
    if strategy is None:
        support_bucket_strategy = None
        return
    val = str(strategy).strip().lower()
    if val not in {"number", "floor"}:
        raise ValueError("support_strategy 必须是 'number' 或 'floor'")
    support_bucket_strategy = val


def set_net_strategy(strategy: str | None):
    """设置全局网架分桶策略。"""
    global net_bucket_strategy
    if strategy is None:
        net_bucket_strategy = None
        return
    val = str(strategy).strip().lower()
    if val not in {"number", "floor"}:
        raise ValueError("net_strategy 必须是 'number' 或 'floor'")
    net_bucket_strategy = val


def merge_remains_into_last_bucket(cats_by_bucket: dict, remain_by_cat: dict):
    """把未分配的数据并入最后一个桶。"""
    if not cats_by_bucket:
        return
    last_idx = None
    for bucket_map in cats_by_bucket.values():
        if bucket_map:
            cur_max = max(bucket_map.keys())
            last_idx = cur_max if last_idx is None else max(last_idx, cur_max)
    if last_idx is None:
        last_idx = 0
    for cat, remain in (remain_by_cat or {}).items():
        bucket_map = cats_by_bucket.setdefault(cat, {})
        if last_idx not in bucket_map:
            bucket_map[last_idx] = []
        bucket_map[last_idx].extend(remain)
        if hasattr(remain, "clear"):
            remain.clear()


def preview_buckets_generic(cat_byb, remain_by_cat, buckets, categories_present):
    """
     预览日期桶分配结果，询问用户是否确认生成，支持未分配数据处理。

     显示每天各类型构件的分配数量及未分配数据；提供选项：
     - 回车：确认生成
     - n：取消操作
     - a：将未分配数据并入最后一天

     Args:
         cat_byb: 按类型和桶分配的结果
         remain_by_cat: 未分配数据
         buckets: 日期桶配置列表
         categories_present: 存在的构件类型列表
     Returns:
         tuple: 包含两个元素的元组：
             - 是否确认生成（bool）
             - 是否将未分配数据并入最后一天（bool）
     """
    print("\n🧾 预览：")
    for i, b in enumerate(buckets, start=1):
        parts = []
        for cat in categories_present:
            parts.append(f"{cat} {len(cat_byb[cat][i - 1])}")
        print(f"  第{i}天 〔{b['date'] or b['date_raw'] or '未填日期'}〕 → " + "、".join(parts))
    if any(remain_by_cat[cat] for cat in categories_present):
        print("  ⚠️ 未分配：", end="")
        print("、".join(f"{cat} {len(remain_by_cat[cat])}" for cat in categories_present if remain_by_cat[cat]))
    ans = ask("确认生成吗？(回车=是 / n=否 / a=把未分配并入最后一天)：", lower=True)
    return (ans != "n"), (ans == "a")


def expand_blocks_by_bucket(cat_byb):
    """
    将按日期桶分配的构件数据组拆分为标准数据块（5行/块）。

    对每个类型、每个日期桶的数据组应用expand_blocks函数，确保数据块结构统一，适配Excel模板。

    Args:
        cat_byb: 按类型和桶分配的结果（assign_by_buckets返回的cat_byb）
    Returns:
        dict: 按类型和桶组织的数据块字典（dict[类型][桶索引] = 数据块列表）
    """
    # 返回：blocks_by_cat[cat][bucket_index] = [blocks...]
    return {cat: {bi: expand_blocks(lst, PER_LINE_PER_BLOCK) for bi, lst in byb.items()}
            for cat, byb in cat_byb.items()}


def ensure_pages_slices_for_cat(wb, cat: str, blocks_by_bucket_for_cat: dict):
    """
    为指定类型的每个日期桶确保足够的工作表，返回按桶划分的工作表切片。

    计算每个桶所需工作表数量（按5块/页），不足时自动复制补充：
    - 常规类型（钢柱/钢梁/支撑）从自身基础表复制
    - “其他”类型从钢柱模板复制
    返回按桶分组的工作表名称列表。

    Args:
        wb: Excel工作簿对象
        cat: 构件类型（如“钢柱”“其他”）
        blocks_by_bucket_for_cat: 该类型按桶组织的数据块字典
    Returns:
        list[list[str]]: 按桶划分的工作表名称列表（每个元素为一个桶的工作表）
    """

    def need_pages(lst):
        return math.ceil(len(lst) / BLOCKS_PER_SHEET) if lst else 0

    page_need_each = [need_pages(blocks_by_bucket_for_cat.get(i, [])) for i in range(len(blocks_by_bucket_for_cat))]
    total_need = sum(page_need_each)
    if total_need == 0:
        return [[] for _ in page_need_each]
    if cat == "其他":
        pages_all = ensure_total_pages_from(wb, "钢柱", "其他", total_need)
    else:
        pages_all = ensure_total_pages(wb, cat, total_need)
    slices = [];
    p = 0
    for n in page_need_each:
        slices.append(pages_all[p:p + n]);
        p += n
    return slices


def make_target_order_generic(pages_slices_by_cat, categories_present):
    """
    生成工作表的目标顺序，按“日期桶→类型优先级”排序。

    排序规则：
    1. 按日期桶轮次分组
    2. 同轮次内按CATEGORY_ORDER（钢柱→钢梁→支撑→网架→其他）排序
    确保工作表按检测流程和类型逻辑有序排列。

    Args:
        pages_slices_by_cat: 按类型和桶划分的工作表切片字典
        categories_present: 存在的构件类型列表
    Returns:
        list[str]: 排序后的工作表名称列表
    """
    rounds = 0
    for cat in categories_present:
        rounds = max(rounds, len(pages_slices_by_cat.get(cat, [])))
    target = []
    for i in range(rounds):
        for cat in CATEGORY_ORDER:
            if cat not in categories_present: continue
            sl = pages_slices_by_cat[cat][i] if i < len(pages_slices_by_cat[cat]) else []
            target += sl
    return target


# ===== Excel 写入带进度 =====
class Prog:
    def __init__(self, total: int, label: str = "写入 Excel"):
        self.total = max(1, total)
        self.done = 0
        self.label = label

    def tick(self, k=1):
        self.done += k
        pct = int(self.done * 100 / self.total)
        sys.stdout.write(f"\r📊 {self.label}：{self.done}/{self.total}（{pct}%）")
        sys.stdout.flush()

    def finish(self):
        sys.stdout.write("\n");
        sys.stdout.flush()


def fill_blocks_to_pages(wb, pages_slice, blocks, prog: Prog | None = None):
    """
    强校验版（页内也拦）：只要发现“当前块”的 类别/μ 与“当前页”不一致，
    - 若当前页还没写（pos==0）：跳过这张页找下一张；
    - 若当前页已写过（pos>0）：先补斜杠收尾，换到下一张再写。
    """
    if not pages_slice:
        return

    max_cap = len(pages_slice) * BLOCKS_PER_SHEET
    if len(blocks) > max_cap:
        sys.stdout.write(f"\n⚠️ 写入块 {len(blocks)} 超出可用容量 {max_cap}（将自动截断，不会串页）。\n")

    page_idx, pos = 0, 0
    i = 0
    while i < len(blocks) and page_idx < len(pages_slice):
        ws = wb[pages_slice[page_idx]]
        title = ws.title

        if STRICT_CROSS_CAT_GUARD:
            sheet_cat = _sheet_cat_from_title(title)
            sheet_is_mu = _is_mu_title(title)
            block_cat = kind_of((blocks[i].get("name") if blocks[i] else "") or "")
            block_is_mu = _is_mu_block(blocks[i])

            mismatch = (sheet_cat and block_cat and sheet_cat != block_cat) or (sheet_is_mu != block_is_mu)
            if mismatch:
                # 页首：这张直接跳过；页中：先收尾再换页
                if pos != 0:
                    slash_tail(ws, detect_anchors(ws), pos)
                page_idx += 1
                pos = 0
                continue

        # 写入当前块
        anc = detect_anchors(ws)
        write_block(ws, anc, pos, blocks[i])
        if prog:
            prog.tick(1)
        pos += 1
        i += 1

        # 换页
        if pos == BLOCKS_PER_SHEET:
            page_idx += 1
            pos = 0

    # 尾页补“/”
    if page_idx < len(pages_slice) and pos != 0:
        ws = wb[pages_slice[page_idx]]
        slash_tail(ws, detect_anchors(ws), pos)




def cleanup_unused_sheets(wb, used_names, bases=("钢柱", "钢梁", "支撑", "网架", "其他")):
    """
    清理Excel中未使用的指定类型工作表，减少冗余。

    仅保留已使用的目标类型工作表（钢柱/钢梁/支撑/网架/其他），避免模板中多余工作表干扰。
    确保至少保留一个工作表（防止工作簿为空）。

    Args:
        wb: Excel工作簿对象
        used_names: 已使用的工作表名称列表
        bases: 目标类型基础名称列表
    """
    # 如果没有任何工作表被使用，则不进行清理，避免误删模板页
    if not used_names:
        return
    used = set(used_names)
    to_remove = []
    for ws in list(wb.worksheets):
        if any(ws.title == b or ws.title.startswith(f"{b}（") for b in bases):
            if ws.title not in used:
                to_remove.append(ws)
    if len(to_remove) >= len(wb.worksheets):
        to_remove = to_remove[:-1]
    for ws in to_remove:
        wb.remove(ws)


def _distribute_by_dates(items, date_entries):
    """按日期列表将项目分配到各天。"""
    res = []
    if not date_entries:
        return res
    if date_entries[0][1] is not None:  # 配额模式
        cursor = 0
        total = len(items)
        n = len(date_entries)
        for i, (d, limit) in enumerate(date_entries):
            remaining = max(0, total - cursor)
            if remaining <= 0:
                res.append((d, []))
                continue
            if limit is None or limit <= 0:
                take = remaining
            elif i < n - 1:
                take = min(int(limit), remaining)
            else:
                take = remaining
            res.append((d, items[cursor:cursor + take]))
            cursor += take
    else:  # 均分
        days = len(date_entries)
        per = math.ceil(len(items) / days) if days else 0
        cursor = 0
        for i, (d, _) in enumerate(date_entries):
            if i < days - 1:
                take = min(per, len(items) - cursor)
            else:
                take = len(items) - cursor
            res.append((d, items[cursor:cursor + take]))
            cursor += take
    return res

def _prompt_dates_and_limits():
    """交互获取日期和每日数量。"""
    while True:
        txt = ask(
            "日期（空格/逗号分隔；支持 20250101 / 2025年1月1日 / 2025 1 1 / 2025.1.1 / 2025-1-1 / 1-1 / 01-01，\n"
            "年份默认取首个日期的年或当前年）：例如 2025-08-27 8-28 2025年1月1日\n→ "
        )
        if any(ch in txt for ch in "；;，、/\\|"):
            print("只接受逗号或空格分隔，请重输。")
            continue
        dates, ig = _parse_dates_simple(txt)
        if not dates:
            print("请输入至少一个合法日期。")
            continue
        if ig:
            print("已忽略：" + "、".join(ig))
        break
    while True:
        txt = ask("每日数量（按日期顺序；空=均分；填整数=配额）\n→ ")
        if txt == "":
            limits = [None] * len(dates)
            break
        tokens = [t for t in re.split(r"[ ,]+", txt) if t]
        if all(t.isdigit() and int(t) > 0 for t in tokens):
            if len(tokens) == 1:
                limits = [int(tokens[0])] * len(dates)
                break
            if len(tokens) == len(dates):
                limits = [int(t) for t in tokens]
                break
        print(f"请输入{len(dates)}个正整数或留空。")

    return list(zip(dates, limits))


def _summarize_plan(tag, plan, all_floors=None):
    """输出楼层计划摘要，便于用户确认。"""

    def fmt(entry):
        ds = " ".join(normalize_date(x[0]) for x in entry)
        ls = ",".join(str(x[1]) if x[1] is not None else "-" for x in entry)
        return f"{ds} → {ls}"

    specified = [f for f in plan if f != "*"]
    if specified:
        print("已单独配置：")
        for f in sorted(specified, key=_floor_sort_key_by_label):
            print(f"  {f} → {fmt(plan[f])}")
    if "*" in plan:
        print("默认配置：")
        print(f"  * → {fmt(plan['*'])}")
    if all_floors:
        miss = [f for f in all_floors if f not in plan and "*" not in plan]
        if miss:
            miss_txt = " ".join(sorted(miss, key=_floor_sort_key_by_label))
            print(f"未覆盖的楼层：{miss_txt} （稍后统一处理/回落到日期分桶）")


def _prompt_plan_for_floors(floors, shared=True):
    """针对给定楼层集合交互生成计划。"""
    floors = sorted(set(floors), key=_floor_sort_key_by_label)
    if floors:
        print("已识别楼层：" + " ".join(floors))
    # Step1 楼层
    while True:
        txt = ask("适用楼层（回车=全部）：示例 5F, 6F, B2, 屋面 或 5 6 B2\n→ ")
        if any(ch in txt for ch in "；;，、/\\|"):
            print("只接受逗号或空格分隔，请重输。")
            continue
        if not txt:
            sel = None
            break
        tokens = [t for t in re.split(r"[ ,]+", txt) if t]
        seen, sel, ig = set(), [], []
        for t in tokens:
            lb = _floor_label_from_name(t)
            if lb != "F?" and lb in floors and lb not in seen:
                sel.append(lb);
                seen.add(lb)
            else:
                ig.append(t)
        if ig:
            print("已忽略：" + "、".join(ig))
        if sel:
            break
        print("没有合法楼层，请重输。")
    targets = floors if sel is None else sel
    if shared:
        print("下面输入的日期与每日上限，将自动应用到以上所有楼层")
        date_entries = _prompt_dates_and_limits()
        if sel is None:
            return {"*": date_entries}
        return {f: date_entries for f in targets}
    plan = {}
    for f in targets:
        print(f"{f}：")
        plan[f] = _prompt_dates_and_limits()
    return plan


def prompt_mode4_plan(floors_by_cat, categories_present):
    """模式4交互，分别为各类别获取楼层计划。"""
    print("各类别楼层：")
    for cat in categories_present:
        fls = sorted(floors_by_cat.get(cat, set()), key=_floor_sort_key_by_label)
        print(f"{cat}: {(' '.join(fls)) if fls else '/'}")
    plans = {}
    for cat in categories_present:
        fls = floors_by_cat.get(cat, set())
        if not fls:
            continue
        print(f"\n[{cat}]")
        share = ask("这些楼层用同一套日期/数量吗？（y=是，回车=分别设置）\n→ ") == "y"
        plans[cat] = _prompt_plan_for_floors(fls, shared=share)
        # —— 新增：给未指定楼层兜底 ——
        all_floors = sorted(floors_by_cat.get(cat, set()), key=_floor_sort_key_by_label)
        plan_for_cat = plans[cat]
        specified = {f for f in plan_for_cat.keys() if f != "*"}
        if "*" not in plan_for_cat and len(specified) < len(all_floors):
            miss = [f for f in all_floors if f not in specified]
            print(f"👉 {cat} 还有未配置楼层：{' '.join(miss)}")
            ans = ask(
                "要不要给“未配置”的楼层用一套通用的日期/数量？（y=是，回车=跳过；未配置的楼层稍后会再统一询问或回落到日期分桶）",
                lower=True
            )
            if ans == "y":
                plan_for_cat["*"] = _prompt_dates_and_limits()
        _summarize_plan(cat, plan_for_cat, all_floors)
    return plans


def mode4_run(wb, grouped, categories_present):
    """执行模式4：按楼层和日期写入Excel。"""
    write_dates = bool(globals().get("NONINTERACTIVE_MODE4_WRITE_DATES", True))
    injected_support = (globals().get("NONINTERACTIVE_MODE4_SUPPORT_STRATEGY") or "").lower()
    if injected_support in {"number", "floor"}:
        set_support_strategy(injected_support)
    injected_net = (globals().get("NONINTERACTIVE_MODE4_NET_STRATEGY") or "").lower()
    if injected_net in {"number", "floor"}:
        set_net_strategy(injected_net)
    cf_groups = defaultdict(list)
    floors_by_cat = defaultdict(set)
    for cat in categories_present:
        for g in grouped[cat]:
            fl = _floor_label_from_name(g["name"])
            cf_groups[(cat, fl)].append(g)
            floors_by_cat[cat].add(fl)
    plan_dict = globals().get("NONINTERACTIVE_MODE4_PLAN")
    if not plan_dict:
        plan_dict = prompt_mode4_plan(floors_by_cat, categories_present)

    blocks_by_cat_bucket = {cat: defaultdict(list) for cat in CATEGORY_ORDER}
    buckets = []  # list[{date}]
    date_idx = {}
    leftover_by_cat = defaultdict(list)

    for (cat, fl), items in cf_groups.items():
        items.sort(key=lambda x: (
            int(re.search(r"\d+", x["name"]).group()) if re.search(r"\d+", x["name"]) else 10 ** 9, x["name"]))
        plan_for_cat = plan_dict.get(cat, {})
        plan = plan_for_cat.get(fl) or plan_for_cat.get("*")
        if not plan:
            leftover_by_cat[cat].extend(items)
            continue
        for date, slice_items in _distribute_by_dates(items, plan):
            if not slice_items:
                continue
            if date not in date_idx:
                date_idx[date] = len(buckets)
                buckets.append({"date": date})
            idx = date_idx[date]
            blocks_by_cat_bucket[cat][idx].extend(expand_blocks(slice_items, PER_LINE_PER_BLOCK))
    # —— 兜底 ——
    left_total = sum(len(v) for v in leftover_by_cat.values())
    if left_total:
        injected_fallback = (globals().get("NONINTERACTIVE_MODE4_FALLBACK") or "").lower()
        default_entries = globals().get("NONINTERACTIVE_MODE4_DEFAULT") or []
        handled_noninteractive = False
        if injected_fallback:
            if injected_fallback == "default":
                if not default_entries:
                    raise RuntimeError("Mode4 fallback=default 但未提供 default_entries")
                for cat in CATEGORY_ORDER:
                    if not leftover_by_cat.get(cat):
                        continue
                    for date, slice_items in _distribute_by_dates(leftover_by_cat[cat], default_entries):
                        if not slice_items:
                            continue
                        if date not in date_idx:
                            date_idx[date] = len(buckets)
                            buckets.append({"date": date})
                        idx = date_idx[date]
                        blocks_by_cat_bucket[cat][idx].extend(expand_blocks(slice_items, PER_LINE_PER_BLOCK))
                    leftover_by_cat[cat] = []
                leftover_by_cat = defaultdict(list)
                handled_noninteractive = True
            elif injected_fallback == "append_last":
                if not buckets:
                    raise RuntimeError("Mode4 append_last 需要至少一个日期桶")
                last_idx = len(buckets) - 1
                for cat in CATEGORY_ORDER:
                    if not leftover_by_cat.get(cat):
                        continue
                    blocks = expand_blocks(leftover_by_cat[cat], PER_LINE_PER_BLOCK)
                    blocks_by_cat_bucket[cat][last_idx].extend(blocks)
                leftover_by_cat = defaultdict(list)
                handled_noninteractive = True
            elif injected_fallback == "error":
                raise RuntimeError("Mode4 未分配楼层未指定处理方案（fallback=error）")
        if not handled_noninteractive:
            print(f"⚠️ 还有 {left_total} 组未分配。")
            ans = ask("是否给未指定楼层套用【默认日期/数量】？(y=是 / 回车=否→回落到日期分桶)", lower=True)
            if ans == "y":
                default_entries = _prompt_dates_and_limits()
                for cat in CATEGORY_ORDER:
                    if not leftover_by_cat.get(cat):
                        continue
                    for date, slice_items in _distribute_by_dates(leftover_by_cat[cat], default_entries):
                        if not slice_items:
                            continue
                        if date not in date_idx:
                            date_idx[date] = len(buckets)
                            buckets.append({"date": date})
                        idx = date_idx[date]
                        blocks_by_cat_bucket[cat][idx].extend(expand_blocks(slice_items, PER_LINE_PER_BLOCK))
                    leftover_by_cat[cat] = []
            else:
                grouped_left = {c: leftover_by_cat[c] for c in CATEGORY_ORDER if leftover_by_cat.get(c)}
                if grouped_left:
                    buckets2 = prompt_date_buckets(list(grouped_left.keys()), grouped_left)
                    later_first = prompt_bucket_priority()
                    cat_byb, remain_by_cat = assign_by_buckets(grouped_left, buckets2, later_first)
                    ok, auto_last = preview_buckets_generic(cat_byb, remain_by_cat, buckets2, list(grouped_left.keys()))
                    if ok:
                        if auto_last:
                            last = len(buckets2) - 1
                            for c in grouped_left.keys():
                                cat_byb[c][last].extend(remain_by_cat[c])
                                remain_by_cat[c] = []
                        blocks_by_cat_bucket2 = expand_blocks_by_bucket(cat_byb)
                        for i, bk in enumerate(buckets2):
                            date = bk["date"]
                            if date not in date_idx:
                                date_idx[date] = len(buckets)
                                buckets.append({"date": date})
                            idx = date_idx[date]
                            for c in grouped_left.keys():
                                blocks_by_cat_bucket[c][idx].extend(blocks_by_cat_bucket2[c].get(i, []))
                        leftover_by_cat = remain_by_cat
                    else:
                        print("❌ 已取消兜底分配。")

    unassigned = sum(len(v) for v in leftover_by_cat.values())

    # —— 日期按升序排序 ——
    order = sorted(range(len(buckets)), key=lambda i: buckets[i]["date"])
    buckets = [buckets[i] for i in order]
    for cat in CATEGORY_ORDER:
        blocks_by_cat_bucket[cat] = {new_i: blocks_by_cat_bucket[cat].get(old_i, []) for new_i, old_i in
                                     enumerate(order)}

    # —— 统一写页 ——
    cats_in_use = [c for c in CATEGORY_ORDER if blocks_by_cat_bucket[c]]
    pages_slices_by_cat = {}
    for cat in cats_in_use:
        blocks_dict = {i: blocks_by_cat_bucket[cat].get(i, []) for i in range(len(buckets))}
        pages_slices_by_cat[cat] = ensure_pages_slices_for_cat(wb, cat, blocks_dict)

    target = make_target_order_generic(pages_slices_by_cat, cats_in_use)
    for idx, name in enumerate(target):
        cur = wb.sheetnames.index(name)
        if cur != idx:
            wb.move_sheet(wb[name], idx - cur)

    total_blocks = 0
    for cat in cats_in_use:
        for i in range(len(buckets)):
            total_blocks += len(blocks_by_cat_bucket[cat].get(i, []))
    prog = Prog(total_blocks, "写入 Excel")
    for i in range(len(buckets)):
        day_pages = []
        for cat in CATEGORY_ORDER:
            if cat not in cats_in_use:
                continue
            pages = pages_slices_by_cat[cat][i]
            blocks = blocks_by_cat_bucket[cat].get(i, [])
            fill_blocks_to_pages(wb, pages, blocks, prog)
            day_pages += pages
        if write_dates:
            apply_meta_on_pages(
                wb,
                day_pages,
                normalize_date(buckets[i]["date"]),
            )
    prog.finish()

    used_names_total = target
    if unassigned:
        print(f"⚠️ 未指派：{unassigned} 组")
    return used_names_total


def try_handle_mode4(mode, wb, grouped, categories_present) -> list | None:
    """模式4兼容钩子。"""
    if mode != "4":
        return None
    return mode4_run(wb, grouped, categories_present)


# ===== 旧法子模式 =====
def prompt_break_submode(has_gz, has_gl):
    """
    交互式选择楼层断点子模式，适配不同数据场景。

    根据是否同时存在钢柱和钢梁提供选项：
    - 同时存在：支持共用断点、分别断点或无断点
    - 仅单类：支持无断点或分别断点
    确保子模式适配实际数据类型。

    Args:
        has_gz: 是否存在钢柱数据（bool）
        has_gl: 是否存在钢梁数据（bool）
    Returns:
        str: 子模式编号（"1"|"2"|"3"）
    """
    if has_gz and has_gl:
        t = ask("断点子模式：1) 柱梁共用断点（简便）  2) 柱梁分别断点  3) 无断点（整单同一天）")
        return t if t in ("1", "2", "3") else "1"
    else:
        t = ask("断点子模式：仅存在单类（或加“其他”） → 3) 无断点  或  2) 分别断点（按各自断点）")
        return t if t in ("2", "3") else "3"


# ===== 主流程 =====
def _parse_breaks_text(text: str) -> list[int]:
    tokens = re.split(r"[\s,，,;；、]+", str(text or ""))
    vals: set[int] = set()
    for tok in tokens:
        tok = tok.strip()
        if not tok:
            continue
        norm = re.sub(r"[~～〜－—–−至到]", "-", tok)
        m_range = re.fullmatch(r"(\d+)\s*-\s*(\d+)", norm)
        if m_range:
            start = int(m_range.group(1))
            end = int(m_range.group(2))
            if start <= end:
                vals.update(range(start, end + 1))
            else:
                vals.update(range(end, start + 1))
            continue
        m_single = re.search(r"(\d+)", tok)
        if not m_single:
            continue
        try:
            vals.add(int(m_single.group(1)))
        except ValueError:
            continue
    return sorted(vals)


def _segment_blocks_by_floor(blocks, breaks: list[int]):
    buckets = defaultdict(list)
    for blk in blocks or []:
        seg = segment_index(floor_of(blk.get("name", "")), breaks)
        buckets[seg].append(blk)
    if not buckets:
        buckets[0] = []
    return buckets


def _segment_blocks_by_number(blocks, breaks: list[int], extractor):
    buckets = defaultdict(list)
    for blk in blocks or []:
        raw = extractor(blk.get("name", ""))
        if raw is None:
            seg = len(breaks) if breaks else 0
        else:
            seg = len(breaks)
            for idx, val in enumerate(breaks):
                if raw <= val:
                    seg = idx
                    break
        buckets[seg].append(blk)
    if not buckets:
        buckets[0] = []
    return buckets


def _run_mode2_auto(
    wb,
    grouped,
    categories_present,
    *,
    breaks_gz: str = "",
    breaks_gl: str = "",
    breaks_support: str = "",
    include_support: bool = True,
):
    categories_present = [c for c in categories_present if grouped.get(c)]
    if not include_support and "支撑" in categories_present:
        categories_present = [c for c in categories_present if c != "支撑"]

    breaks_gz_list = _parse_breaks_text(breaks_gz)
    breaks_gl_list = _parse_breaks_text(breaks_gl)
    anchor_breaks = sorted(set(breaks_gz_list + breaks_gl_list))

    sup_breaks_raw = (globals().get("NONINTERACTIVE_SUPPORT_BREAKS") or "").strip()
    sup_breaks_list = _parse_breaks_text(sup_breaks_raw) if sup_breaks_raw else anchor_breaks

    support_strategy = (globals().get("NONINTERACTIVE_SUPPORT_STRATEGY") or "number").lower()
    net_strategy = (globals().get("NONINTERACTIVE_NET_STRATEGY") or "number").lower()
    net_breaks_raw = (globals().get("NONINTERACTIVE_NET_BREAKS") or "").strip()
    net_breaks_list = _parse_breaks_text(net_breaks_raw) if net_breaks_raw else anchor_breaks

    blocks_by_cat = {cat: expand_blocks(grouped.get(cat, []), PER_LINE_PER_BLOCK)
                     for cat in categories_present}

    buckets_by_cat = {}
    segment_ids = set()

    for cat in categories_present:
        blocks = blocks_by_cat.get(cat, [])
        if cat in ("钢柱", "钢梁"):
            buckets = _segment_blocks_by_floor(blocks, anchor_breaks)
        elif cat == "支撑":
            if support_strategy == "floor":
                buckets = _segment_blocks_by_floor(blocks, sup_breaks_list)
            else:
                buckets = _segment_blocks_by_number(blocks, sup_breaks_list, _wz_no)
        elif cat == "网架":
            if net_strategy == "floor":
                buckets = _segment_blocks_by_floor(blocks, net_breaks_list)
            else:
                buckets = _segment_blocks_by_number(blocks, net_breaks_list, _net_no)
        else:
            buckets = defaultdict(list)
            buckets[0] = list(blocks)
        buckets_by_cat[cat] = buckets
        segment_ids.update(buckets.keys())

    if not segment_ids:
        segment_ids = {0}
    ordered_segments = sorted(segment_ids)

    pages_slices_by_cat = {}
    blocks_slices_by_cat = {}

    for cat in categories_present:
        bucket_map = {seg: list(buckets_by_cat[cat].get(seg, [])) for seg in ordered_segments}
        if cat == "其他":
            pages_list = []
            blocks_list = []
            for seg in ordered_segments:
                seg_blocks = bucket_map.get(seg, [])
                need = pages_needed(seg_blocks)
                pages_batch = [] if not need else ensure_total_pages_from(wb, "钢柱", "其他", need)
                pages_list.append(pages_batch)
                blocks_list.append(seg_blocks)
            pages_slices_by_cat[cat] = pages_list
            blocks_slices_by_cat[cat] = blocks_list
        else:
            pages_slices, blocks_slices = ensure_pages_slices_for_cat_muaware(wb, cat, bucket_map)
            pages_slices_by_cat[cat] = [_filter_pages_for_cat(sl, cat) for sl in pages_slices]
            blocks_slices_by_cat[cat] = blocks_slices

    total_blocks = 0
    for cat in categories_present:
        for seg_blocks in blocks_slices_by_cat[cat]:
            total_blocks += len(seg_blocks)

    prog = Prog(total_blocks or 1, "写入 Excel")

    used_pages: list[str] = []
    date_first = (globals().get("NONINTERACTIVE_MODE2_DATE_FIRST") or "").strip()
    date_second = (globals().get("NONINTERACTIVE_MODE2_DATE_SECOND") or "").strip()
    norm_first = normalize_date(date_first) if date_first else ""
    norm_second = normalize_date(date_second) if date_second else ""

    for seg_idx, _seg in enumerate(ordered_segments):
        for cat in CATEGORY_ORDER:
            if cat not in categories_present:
                continue
            pages_list = pages_slices_by_cat.get(cat, [])
            blocks_list = blocks_slices_by_cat.get(cat, [])
            if seg_idx >= len(pages_list):
                continue
            pages = pages_list[seg_idx]
            blocks_piece = blocks_list[seg_idx]
            if not pages:
                continue
            fill_blocks_to_pages(wb, pages, blocks_piece, prog)
            used_pages.extend(pages)
            date_to_write = norm_first if seg_idx == 0 else (norm_second or norm_first)
            if date_to_write:
                apply_meta_on_pages(wb, pages, date_to_write)

    prog.finish()

    for idx, name in enumerate(used_pages):
        if name not in wb.sheetnames:
            continue
        cur = wb.sheetnames.index(name)
        if cur != idx:
            wb.move_sheet(wb[name], idx - cur)

    cleanup_unused_mu_templates(wb, used_pages)
    return used_pages


def run_mode(
    mode: str,
    wb,
    grouped=None,
    categories_present=None,
    *,
    src: Union[str, Path] | None = None,
    grouped_preloaded=None,
    breaks_gz: str = "",
    breaks_gl: str = "",
    include_support: bool = True,
):
    """按指定模式执行一次导出（全模式支持 μ 逻辑；mode4 暂保持原样流程）。"""
    global support_bucket_strategy, net_bucket_strategy
    support_bucket_strategy = None
    net_bucket_strategy = None

    if grouped_preloaded is not None:
        grouped_data = grouped_preloaded
    elif grouped is not None:
        grouped_data = grouped
    elif src is not None:
        grouped_data, categories_from_src = prepare_from_word(Path(src))
        if categories_present is None:
            categories_present = categories_from_src
    elif _PROBE_CACHE.get("src") and Path(str(_PROBE_CACHE.get("src"))).exists() and src is None:
        grouped_data = _PROBE_CACHE.get("grouped") or {}
    else:
        raise ValueError("run_mode 需要提供 grouped/grouped_preloaded/src 之一")

    if isinstance(grouped_data, dict) and not isinstance(grouped_data, defaultdict):
        tmp = defaultdict(list)
        for k, v in grouped_data.items():
            tmp[k] = list(v)
        grouped_data = tmp

    if categories_present is None:
        categories_present = [cat for cat in CATEGORY_ORDER if grouped_data.get(cat)]

    # 先交给 mode4 的专用处理（不动它内部逻辑）
    res = try_handle_mode4(mode, wb, grouped_data, categories_present)
    if res is not None:
        return res

    force_same_breaks = bool(globals().get("NONINTERACTIVE_MODE2_FORCE_SAME_BREAKS"))
    if mode == "2" and (grouped_preloaded is not None or force_same_breaks):
        return _run_mode2_auto(
            wb,
            grouped_data,
            categories_present,
            breaks_gz=breaks_gz,
            breaks_gl=breaks_gl,
            include_support=include_support,
        )

    # ============ mode 2：按楼层断点 ============
    if mode == "2":
        has_gz = "钢柱" in categories_present
        has_gl = "钢梁" in categories_present
        sub = prompt_break_submode(has_gz, has_gl)

        blocks_by_cat = {cat: expand_blocks(grouped_data[cat], PER_LINE_PER_BLOCK)
                         for cat in categories_present}

        # —— 子模式 3：无断点，整类一次性排（也用 μ-aware）——
        if sub == "3":
            pages_by_cat = {}
            blocks_by_cat_ordered = {}

            for cat in categories_present:
                blocks_all = blocks_by_cat[cat]
                if cat == "其他":
                    need = pages_needed(blocks_all)
                    pages_by_cat[cat] = [] if not need else ensure_total_pages_from(wb, "钢柱", "其他", need)
                    blocks_by_cat_ordered[cat] = blocks_all
                else:
                    # 复用 μ-aware，视作“只有一个桶”，索引 0
                    pages_slices, blocks_slices = ensure_pages_slices_for_cat_muaware(
                        wb, cat, {0: blocks_all}
                    )
                    pages_by_cat[cat] = pages_slices[0]
                    blocks_by_cat_ordered[cat] = blocks_slices[0]

            target = []
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    target += pages_by_cat[cat]
            for idx, name in enumerate(target):
                cur = wb.sheetnames.index(name)
                if cur != idx:
                    wb.move_sheet(wb[name], idx - cur)

            total_blocks = sum(len(blocks_by_cat_ordered[cat]) for cat in categories_present)
            prog = Prog(total_blocks, "写入 Excel")
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    fill_blocks_to_pages(wb, pages_by_cat[cat], blocks_by_cat_ordered[cat], prog)
            prog.finish()

            d = normalize_date(ask("📅 整单日期（回车=不写）：") or "")
            apply_meta_on_pages(wb, target, d)
            cleanup_unused_mu_templates(wb, target)
            return target

        # —— 子模式 1/2：按断点分段（每段也是 μ-aware）——
        same_breaks = None
        if has_gz and has_gl and sub == "1":
            same_breaks = prompt_floor_breaks("钢柱/钢梁（共用）")

        breaks_by_cat = {}
        for cat in categories_present:
            if cat == "支撑":
                prompt_support_strategy_for_bucket()
                if support_bucket_strategy == "floor":
                    breaks_by_cat[cat] = prompt_floor_breaks(cat)
                else:
                    breaks_by_cat[cat] = []  # 支撑不分段
            elif cat in ("钢柱", "钢梁"):
                if ((cat == "钢柱" and "钢梁" in categories_present) or
                    (cat == "钢梁" and "钢柱" in categories_present)) and same_breaks is not None:
                    breaks_by_cat[cat] = same_breaks
                else:
                    breaks_by_cat[cat] = prompt_floor_breaks(cat)
            else:
                breaks_by_cat[cat] = prompt_floor_breaks(cat)

        # 建段：用 floor_of + segment_index
        byseg = {cat: defaultdict(list) for cat in categories_present}
        for cat in categories_present:
            if cat == "支撑" and support_bucket_strategy != "floor":
                byseg[cat][0] = blocks_by_cat[cat]
            else:
                for b in blocks_by_cat[cat]:
                    seg = segment_index(floor_of(b["name"]), breaks_by_cat[cat])
                    byseg[cat][seg].append(b)

        # 先对每个类别一次性切片，保证编号连续
        pages_slices_by_cat = {}
        blocks_slices_by_cat = {}
        for cat in categories_present:
            seg_dict = byseg[cat]
            if cat == "其他":
                pages_slices_by_cat[cat] = []
                blocks_slices_by_cat[cat] = []
                for seg in sorted(seg_dict.keys()):
                    seg_blocks = seg_dict[seg]
                    need = pages_needed(seg_blocks)
                    pages_batch = [] if not need else ensure_total_pages_from(wb, "钢柱", "其他", need)
                    pages_slices_by_cat[cat].append(pages_batch)
                    blocks_slices_by_cat[cat].append(seg_blocks)
            else:
                pages_slices_by_cat[cat], blocks_slices_by_cat[cat] = ensure_pages_slices_for_cat_muaware(
                    wb, cat, seg_dict
                )
                pages_slices_by_cat[cat] = [_filter_pages_for_cat(sl, cat) for sl in pages_slices_by_cat[cat]]

        # 构造 (pages, blocks) 队列，按 类×段 逐对写入
        rounds = max(len(pages_slices_by_cat[c]) for c in categories_present)
        pairs = []
        for i in range(rounds):
            for cat in CATEGORY_ORDER:
                if cat not in categories_present:
                    continue
                p_list = pages_slices_by_cat[cat]
                b_list = blocks_slices_by_cat[cat]
                if i < len(p_list) and p_list[i]:
                    pairs.append((p_list[i], b_list[i]))

        target = []
        prog = Prog(sum(len(b) for _, b in pairs), "写入 Excel")
        for pages, blocks_piece in pairs:
            target += pages
            fill_blocks_to_pages(wb, pages, blocks_piece, prog)
        prog.finish()

        # 调整顺序并写入元信息
        for idx, name in enumerate(target):
            cur = wb.sheetnames.index(name)
            if cur != idx:
                wb.move_sheet(wb[name], idx - cur)

        apply_meta_on_pages(wb, target, "")
        cleanup_unused_mu_templates(wb, target)
        return target

    # ============ mode 3：单日模式（已有 μ 逻辑，这里接到 μ-aware） ============
    elif mode == "3":
        pages_by_cat = {}
        blocks_by_cat_ordered = {}

        for cat in categories_present:
            blocks_all = expand_blocks(grouped_data[cat], PER_LINE_PER_BLOCK)
            if cat == "其他":
                need = pages_needed(blocks_all)
                pages_by_cat[cat] = [] if not need else ensure_total_pages_from(wb, "钢柱", "其他", need)
                blocks_by_cat_ordered[cat] = blocks_all
            else:
                pages_slices, blocks_slices = ensure_pages_slices_for_cat_muaware(
                    wb, cat, {0: blocks_all}
                )
                pages_by_cat[cat] = _filter_pages_for_cat(pages_slices[0], cat)  # 👈 新增过滤
                blocks_by_cat_ordered[cat] = blocks_slices[0]

        target = []
        for cat in CATEGORY_ORDER:
            if cat in categories_present:
                target += pages_by_cat[cat]
        for idx, name in enumerate(target):
            cur = wb.sheetnames.index(name)
            if cur != idx:
                wb.move_sheet(wb[name], idx - cur)

        prog = Prog(sum(len(blocks_by_cat_ordered[c]) for c in categories_present), "写入 Excel")
        for cat in CATEGORY_ORDER:
            if cat in categories_present:
                fill_blocks_to_pages(wb, pages_by_cat[cat], blocks_by_cat_ordered[cat], prog)
        prog.finish()

        # 新版：优先使用“非交互注入”的日期，避免 ask 卡住
        _injected = globals().pop("NONINTERACTIVE_MODE3_DATE",
                                  None) if "NONINTERACTIVE_MODE3_DATE" in globals() else None

        if _injected is not None:
            # UI/非交互调用：传 None/"" 表示跳过写日期
            _date_in = _injected
        else:
            # 仅在 CLI 交互时才询问
            try:
                _date_in = ask("📅 请输入检测日期（回车跳过；输入 q 返回上一步）：")
            except BackStep:
                raise

        if str(_date_in).strip():
            apply_meta_on_pages(wb, target, normalize_date(str(_date_in)))
        else:
            apply_meta_on_pages(wb, target, "")

        cleanup_unused_mu_templates(wb, target)
        return target


    # ============ mode 1：日期分桶（每个“日桶”也 μ-aware） ============
    elif mode == "1":
        buckets = prompt_date_buckets(categories_present, grouped_data)
        if buckets is None:
            return

        later_first = prompt_later_priority()
        cat_byb, remain_by_cat = assign_by_buckets(grouped_data, buckets, later_first)
        ok, auto_last_preview = preview_buckets_generic(cat_byb, remain_by_cat, buckets, categories_present)
        if not ok:
            return

        forced_choice = prompt_auto_merge_remains(
            remain_by_cat=remain_by_cat,
            buckets=buckets,
            categories_present=categories_present,
            preview_choice=auto_last_preview,
        )
        auto_last = bool(auto_last_preview)
        forced_provided = forced_choice is not None
        if forced_provided:
            auto_last = bool(forced_choice)

        unassigned = sum(len(v) for v in remain_by_cat.values())
        if unassigned and not auto_last and not forced_provided:
            print(f"⚠️ 未指派：{unassigned} 组")
            auto = ask("是否自动把未指派并入最后一天？（y=是 / 其它=否）", allow_empty=False, lower=True)
            if auto == "y":
                auto_last = True
            elif auto == "q":
                raise BackStep()

        if auto_last:
            last = len(buckets) - 1
            for cat in categories_present:
                cat_byb[cat][last].extend(remain_by_cat[cat])
                remain_by_cat[cat] = []

        blocks_by_cat_bucket = expand_blocks_by_bucket(cat_byb)

        # —— 关键：把“每天”的块做成 μ-aware 的切片 ——
        pages_slices_by_cat = {}
        blocks_slices_by_cat = {}
        for cat in categories_present:
            # blocks_by_cat_bucket[cat] 是 dict: day_idx -> blocks(list)
            pages_slices_by_cat[cat], blocks_slices_by_cat[cat] = ensure_pages_slices_for_cat_muaware(
                wb, cat, blocks_by_cat_bucket[cat]
            )
            pages_slices_by_cat[cat] = [_filter_pages_for_cat(sl, cat) for sl in pages_slices_by_cat[cat]]

        # 拼成最终顺序（按天交错：柱→梁→支撑→其他）
        target = []
        num_days = len(buckets)
        for i in range(num_days):
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    target += pages_slices_by_cat[cat][i]

        for idx, name in enumerate(target):
            cur = wb.sheetnames.index(name)
            if cur != idx:
                wb.move_sheet(wb[name], idx - cur)

        # 写入（逐天）
        total_blocks = 0
        for cat in categories_present:
            total_blocks += sum(len(v) for v in blocks_by_cat_bucket[cat].values())
        prog = Prog(total_blocks, "写入 Excel")

        for i in range(num_days):
            day_pages = []
            day_blocks = []
            for cat in CATEGORY_ORDER:
                if cat in categories_present:
                    day_pages += pages_slices_by_cat[cat][i]
                    day_blocks += blocks_slices_by_cat[cat][i]
            fill_blocks_to_pages(wb, day_pages, day_blocks, prog)
            apply_meta_on_pages(wb, day_pages, buckets[i]["date"])

        prog.finish()
        cleanup_unused_mu_templates(wb, target)
        return target

    else:
        raise ValueError(f"未知的模式：{mode}")


    # ===== 预处理与模式运行封装 =====


def prepare_from_word(src: Path):
    groups_all_tables, all_rows = read_groups_from_doc(src)
    grouped = defaultdict(list)
    for g in groups_all_tables:
        grouped[kind_of(g["name"])].append(g)
    categories_present = [cat for cat in CATEGORY_ORDER if grouped.get(cat)]
    print("📊 识别： " + "、".join(f"{cat} {len(grouped.get(cat, []))}" for cat in categories_present))

    doc_out = build_summary_doc_with_progress(all_rows)
    set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
    out_docx = src.with_name("汇总原始记录.docx")
    print("💾 正在保存汇总 Word …")

    save_docx_safe(doc_out, out_docx)
    print(f"✅ 汇总 Word 已保存：{out_docx}")
    return grouped, categories_present


def run_with_mode(src: Path, grouped, categories_present, meta):
    tpl_path = XLSX_WITH_SUPPORT_DEFAULT  # 始终使用有支撑模板
    if not tpl_path.exists():
        raise FileNotFoundError(f"Excel 模板不存在：{tpl_path}")

    wb = load_workbook_safe(tpl_path)

    try:
        mode = prompt_mode()
        used_names_total = run_mode(mode, wb, grouped, categories_present)
    except BackStep:
        return

    apply_meta_fixed(wb, categories_present, meta)
    enforce_mu_font(wb)
    cleanup_unused_sheets(wb, used_names_total, bases=tuple(CATEGORY_ORDER))

    def unique_out_path(dest_dir: Path, stem: str) -> Path:
        cand = dest_dir / f"{stem}.xlsx"
        if not cand.exists():
            return cand
        i = 1
        while True:
            cand = dest_dir / f"{stem}({i}).xlsx"
            if not cand.exists():
                return cand
            i += 1

    final_path = unique_out_path(src.parent, f"{TITLE}_报告版")
    save_workbook_safe(wb, final_path)
    print(f"✅ Excel 已保存：{final_path}")
    print("✔ 完成。本次导出结束。")

# ===== 非交互入口（供 GUI 调用 / 可脚本化） =====
# ====== 日期填充工具（新增） ======
import re
from datetime import datetime
from pathlib import Path

def _normalize_date(date_str: str) -> str:
    """
    接受 '2025-10-13' / '2025/10/13' / '2025.10.13' / '2025年10月13日' / '2025 10 13'
    统一规范为 'YYYY-MM-DD'；不合法则抛异常。
    """
    s = str(date_str).strip()
    if not s:
        raise ValueError("检测日期为空")
    nums = list(map(int, re.findall(r"\d+", s)))
    if len(nums) >= 3:
        y, m, d = nums[:3]
        dt = datetime(year=y, month=m, day=d)
        return dt.strftime("%Y-%m-%d")
    try:
        return datetime.fromisoformat(s).strftime("%Y-%m-%d")
    except Exception:
        raise ValueError(f"无法识别的日期格式：{s}")

def _fill_date_in_sheet(ws, date_text: str) -> bool:
    """
    在单个工作表里寻找“日期/检验日期/探伤日期”字样（前20行×前20列），
    优先写到右侧单元格；若右侧不可写，则把当前单元格文本替换为“……：YYYY-MM-DD”。
    返回是否写入成功。
    """
    ROW_MAX, COL_MAX = 20, 20
    for r in range(1, min(ws.max_row, ROW_MAX) + 1):
        for c in range(1, min(ws.max_column, COL_MAX) + 1):
            cell = ws.cell(r, c)
            v = cell.value
            if isinstance(v, str) and ("日期" in v or "检验日期" in v or "探伤日期" in v):
                # 1) 右侧邻格优先
                try:
                    neighbor = ws.cell(r, c + 1)
                    if neighbor.value in (None, "", "——", "-", "—"):
                        neighbor.value = date_text
                        return True
                except Exception:
                    pass
                # 2) 改当前格文本
                txt = v
                txt = re.sub(r"(检验日期|探伤日期|日期)[:：]?\s*$", r"\1：" + date_text, txt)
                cell.value = txt
                return True
    return False

def apply_date_to_workbook(wb, date_text: str) -> int:
    """把日期写入工作簿的可见工作表；返回成功写入的表数量。"""
    ok = 0
    for ws in wb.worksheets:
        try:
            if _fill_date_in_sheet(ws, date_text):
                ok += 1
        except Exception:
            pass
    return ok


# ====== 非交互入口（替换为这个完整体） ======
def run_noninteractive(
    src_path,
    mode=3,
    meta=None,
    support_strategy=None,   # "number" | "floor"
    net_strategy=None,       # "number" | "floor"
    dates=None,              # 预留：mode1 用
    temperature=None,        # 预留
    quota_plan=None,         # 预留：mode4 用
    single_date=None,        # 新增：单日模式的“检测日期”
):
    """
    一次性执行完整流程（读取 Word → 生成 Excel → 保存），不依赖 input()。
    目前稳定支持 mode=3（单日模式）直跑；其它模式会自动回退至 3，避免卡住。
    返回：{"excel": Path, "word": Path}
    """
    # 1) 校验源
    src = Path(str(src_path)).expanduser().resolve()
    if not src.exists():
        raise FileNotFoundError(f"找不到源文件：{src}")
    if src.suffix.lower() != ".docx":
        raise ValueError("源文件必须为 .docx")

    # 2) 从 Word 读取、分组 & 汇总
    grouped, categories_present = prepare_from_word(src)

    # 3) 设置分桶策略（若传入则覆盖全局）
    global support_bucket_strategy, net_bucket_strategy
    if support_strategy in ("number", "floor"):
        support_bucket_strategy = support_strategy
    if net_strategy in ("number", "floor"):
        net_bucket_strategy = net_strategy

    # 4) 选择模板并载入
    mode_str = str(mode) if str(mode) in {"1", "2", "3", "4"} else "3"
    if mode_str != "3":
        # 当前仅保证单日模式无交互直跑，其它模式回退到 3
        mode_str = "3"

    tpl_path = XLSX_WITH_SUPPORT_DEFAULT
    if not tpl_path.exists():
        raise FileNotFoundError(f"Excel 模板不存在：{tpl_path}")
    wb = load_workbook_safe(tpl_path)

    # 5) 生成填表（按你的内部实现，这里是你已有的“单日模式”入口）
    #    注意：如果你项目里对应函数名是 run_with_mode(...)，请据实替换这一行。
    used_names_total = run_mode(mode_str, wb, grouped, categories_present)

    # 6) 写元信息 & 统一字体 & 清除无用表
    meta = meta or {}
    apply_meta_fixed(wb, categories_present, meta)
    enforce_mu_font(wb)
    cleanup_unused_sheets(wb, used_names_total, bases=tuple(CATEGORY_ORDER))

    # 7) 若传入“检测日期”，规范化并写入工作簿
    if single_date:
        dt_norm = _normalize_date(single_date)
        _ = apply_date_to_workbook(wb, dt_norm)

    # 8) 生成不覆盖的输出路径并保存
    def _unique_out_path(dest_dir: Path, stem: str) -> Path:
        cand = dest_dir / f"{stem}.xlsx"
        if not cand.exists():
            return cand
        i = 1
        while True:
            cand = dest_dir / f"{stem}({i}).xlsx"
            if not cand.exists():
                return cand
            i += 1

    final_xlsx = _unique_out_path(src.parent, f"{TITLE}_报告版")
    save_workbook_safe(wb, final_xlsx)

    # 9) 返回路径
    word_out = src.with_name("汇总原始记录.docx")
    return {"excel": final_xlsx, "word": word_out}


def _norm_entry_list(entries):
    """规范化计划条目列表（日期统一为 YYYY-MM-DD，数量转 int/None）。"""
    out = []
    if not entries:
        return out
    for d, lim in entries:
        nd = None
        last_err = None
        for fn in (normalize_date, _normalize_date):
            if not fn:
                continue
            try:
                nd = fn(d)
                break
            except Exception as exc:  # noqa: PERF203 - 需要逐个尝试
                last_err = exc
        if not nd:
            raise ValueError(f"无法识别的日期：{d}") from last_err
        if lim in (None, "", "-", "∞"):
            nl = None
        else:
            try:
                nl = int(lim)
            except Exception:
                digits = re.findall(r"\d+", str(lim))
                nl = int(digits[0]) if digits else None
        out.append((nd, nl))
    return out


def _norm_plan(plan: dict | None) -> dict:
    """规范化按类别/楼层的计划结构。"""
    if not plan:
        return {}
    result: dict = {}
    for cat, by_floor in plan.items():
        result[cat] = {}
        for floor, entries in (by_floor or {}).items():
            result[cat][floor] = _norm_entry_list(entries)
    return result


def export_mode4_noninteractive(
        src_docx: Union[str, Path],
        meta: dict | None = None,
        wb=None,
        *,
        plan: dict | None = None,
        include_support: bool = True,
        support_strategy: str = "number",
        net_strategy: str = "number",
        fallback: str = "append_last",
        default_entries: list[tuple[str, int | None]] | None = None,
        write_dates_to_header: bool = True,
) -> tuple[Path, Path | None]:
    """无交互导出 Mode4。"""

    src = Path(str(src_docx)).resolve()
    if not src.exists():
        raise FileNotFoundError(f"未找到 Word 源文件：{src}")

    grouped = None
    categories_present = None
    cache_src = _PROBE_CACHE.get("src")
    if cache_src and Path(str(cache_src)).resolve() == src:
        grouped = _PROBE_CACHE.get("grouped") or defaultdict(list)
        categories_present = list(_PROBE_CACHE.get("categories") or [])

    if grouped is None or categories_present is None:
        grouped, categories_present = prepare_from_word(src)
    else:
        if not isinstance(grouped, defaultdict):
            tmp = defaultdict(list)
            for k, v in (grouped or {}).items():
                tmp[k] = list(v)
            grouped = tmp

    categories_present = [cat for cat in CATEGORY_ORDER if grouped.get(cat)]
    if not include_support and "支撑" in categories_present:
        categories_present = [c for c in categories_present if c != "支撑"]

    prev_support = support_bucket_strategy
    prev_net = net_bucket_strategy
    sup_val = (support_strategy or "number").lower()
    net_val = (net_strategy or "number").lower()
    set_support_strategy(sup_val)
    set_net_strategy(net_val)

    if wb is None:
        if not XLSX_WITH_SUPPORT_DEFAULT.exists():
            raise FileNotFoundError(f"Excel 模板不存在：{XLSX_WITH_SUPPORT_DEFAULT}")
        wb = load_workbook_safe(XLSX_WITH_SUPPORT_DEFAULT)

    globals()["NONINTERACTIVE_MODE4_PLAN"] = _norm_plan(plan)
    globals()["NONINTERACTIVE_MODE4_FALLBACK"] = (fallback or "").lower()
    globals()["NONINTERACTIVE_MODE4_DEFAULT"] = _norm_entry_list(default_entries or [])
    globals()["NONINTERACTIVE_MODE4_SUPPORT_STRATEGY"] = sup_val
    globals()["NONINTERACTIVE_MODE4_NET_STRATEGY"] = net_val
    globals()["NONINTERACTIVE_MODE4_WRITE_DATES"] = bool(write_dates_to_header)

    try:
        used_pages = run_mode("4", wb, grouped, categories_present)
        apply_meta_fixed(wb, categories_present, meta or {})
        cleanup_unused_mu_templates(wb, used_pages)

        def _unique_out_path(dest_dir: Path, stem: str) -> Path:
            cand = dest_dir / f"{stem}.xlsx"
            if not cand.exists():
                return cand
            i = 1
            while True:
                cand = dest_dir / f"{stem}({i}).xlsx"
                if not cand.exists():
                    return cand
                i += 1

        final_xlsx = _unique_out_path(src.parent, f"{TITLE}_报告版")
        save_workbook_safe(wb, final_xlsx)

        word_out = src.with_name("汇总原始记录.docx")
        if not word_out.exists():
            all_rows = _PROBE_CACHE.get("all_rows")
            if all_rows:
                try:
                    doc_out = build_summary_doc_with_progress(all_rows)
                    set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
                    save_docx_safe(doc_out, word_out)
                except Exception:
                    word_out = None
            else:
                word_out = None

        return final_xlsx, word_out
    finally:
        set_support_strategy(prev_support)
        set_net_strategy(prev_net)
        for key in (
                "NONINTERACTIVE_MODE4_PLAN",
                "NONINTERACTIVE_MODE4_FALLBACK",
                "NONINTERACTIVE_MODE4_DEFAULT",
                "NONINTERACTIVE_MODE4_SUPPORT_STRATEGY",
                "NONINTERACTIVE_MODE4_NET_STRATEGY",
                "NONINTERACTIVE_MODE4_WRITE_DATES",
        ):
            globals().pop(key, None)


def export_mode1_noninteractive(
        src_docx,
        out_dir=None,
        *,
        buckets,
        support_strategy="number",
        net_strategy="number",
        later_priority=True,
        auto_merge_rest=True,
        meta=None,
):
    """纯无交互导出 Mode1。"""

    provider = Mode1ConfigProvider(
        buckets,
        support_strategy,
        net_strategy,
        later_priority,
        auto_merge_rest,
        meta=meta,
    )

    src = Path(str(src_docx)).resolve()
    if not src.exists():
        raise FileNotFoundError(f"未找到 Word 源文件：{src}")

    out_dir = Path(out_dir) if out_dir is not None else src.parent
    out_dir = out_dir.resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    cache_src = _PROBE_CACHE.get("src")
    grouped = None
    categories_present = None
    if cache_src and Path(str(cache_src)).resolve() == src:
        grouped = _PROBE_CACHE.get("grouped") or defaultdict(list)
        categories_present = list(_PROBE_CACHE.get("categories") or [])
    if grouped is None or categories_present is None:
        grouped, categories_present = prepare_from_word(src)
    else:
        if not isinstance(grouped, defaultdict):
            tmp = defaultdict(list)
            for k, v in (grouped or {}).items():
                tmp[k] = list(v)
            grouped = tmp

    categories_present = [cat for cat in CATEGORY_ORDER if grouped.get(cat)]

    tpl_path = XLSX_WITH_SUPPORT_DEFAULT
    if not tpl_path.exists():
        raise FileNotFoundError(f"Excel 模板不存在：{tpl_path}")
    wb = load_workbook_safe(tpl_path)

    prev_support = support_bucket_strategy
    prev_net = net_bucket_strategy
    set_support_strategy(provider.support_strategy)
    set_net_strategy(provider.net_strategy)

    buckets_norm = provider.get_buckets()
    try:
        cat_byb, remain_by_cat = assign_by_buckets(grouped, buckets_norm, provider.later_priority)
    finally:
        set_support_strategy(prev_support)
        set_net_strategy(prev_net)

    if provider.auto_merge_rest:
        merge_remains_into_last_bucket(cat_byb, remain_by_cat)

    blocks_by_cat_bucket = expand_blocks_by_bucket(cat_byb)

    pages_slices_by_cat = {}
    blocks_slices_by_cat = {}
    for cat in categories_present:
        bucket_map = blocks_by_cat_bucket.get(cat, {})
        pages_slices, blocks_slices = ensure_pages_slices_for_cat_muaware(wb, cat, bucket_map)
        pages_slices_by_cat[cat] = [_filter_pages_for_cat(sl, cat) for sl in pages_slices]
        blocks_slices_by_cat[cat] = blocks_slices

    num_days = len(buckets_norm)
    target = []
    for i in range(num_days):
        for cat in CATEGORY_ORDER:
            if cat in categories_present:
                target += pages_slices_by_cat[cat][i]

    for idx, name in enumerate(target):
        cur = wb.sheetnames.index(name)
        if cur != idx:
            wb.move_sheet(wb[name], idx - cur)

    total_blocks = 0
    for cat in categories_present:
        total_blocks += sum(len(v) for v in blocks_by_cat_bucket.get(cat, {}).values())
    prog = Prog(total_blocks, "写入 Excel")

    for day_idx in range(num_days):
        day_pages = []
        day_blocks = []
        for cat in CATEGORY_ORDER:
            if cat in categories_present:
                day_pages += pages_slices_by_cat[cat][day_idx]
                day_blocks += blocks_slices_by_cat[cat][day_idx]
        fill_blocks_to_pages(wb, day_pages, day_blocks, prog)
        raw = buckets_norm[day_idx].get("date_raw") or buckets_norm[day_idx].get("date") or ""
        dt = normalize_date(raw) if raw else ""
        apply_meta_on_pages(wb, day_pages, dt)

    prog.finish()

    cleanup_unused_mu_templates(wb, target)
    apply_meta_fixed(wb, categories_present, provider.meta)
    enforce_mu_font(wb)
    cleanup_unused_sheets(wb, target, bases=tuple(CATEGORY_ORDER))

    def _unique_out_path(dest_dir: Path, stem: str) -> Path:
        cand = dest_dir / f"{stem}.xlsx"
        if not cand.exists():
            return cand
        i = 1
        while True:
            cand = dest_dir / f"{stem}({i}).xlsx"
            if not cand.exists():
                return cand
            i += 1

    final_xlsx = _unique_out_path(out_dir, f"{TITLE}_报告版")
    save_workbook_safe(wb, final_xlsx)

    word_out = src.with_name("汇总原始记录.docx")
    if not word_out.exists():
        all_rows = _PROBE_CACHE.get("all_rows")
        if all_rows:
            try:
                doc_out = build_summary_doc_with_progress(all_rows)
                set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
                save_docx_safe(doc_out, word_out)
            except Exception:
                pass

    return final_xlsx, word_out


# ===== 非交互：按楼层断点（Mode 2）导出 =====

def export_mode2_noninteractive(
    src_docx: Union[str, Path],
    meta: dict | None = None,
    wb=None,
    *,
    breaks_gz: str = "",
    breaks_gl: str = "",
    breaks_support: str = "",
    breaks_net: str = "",
    date_first: str = "",
    date_second: str = "",
    include_support: bool = True,
    support_strategy: str = "number",
    net_strategy: str = "number",
):
    src = Path(str(src_docx)).resolve()
    if not src.exists():
        raise FileNotFoundError(f"未找到 Word 源文件：{src}")

    grouped = None
    categories_present = None
    cache_src = _PROBE_CACHE.get("src")
    if cache_src and Path(str(cache_src)).resolve() == src:
        grouped = _PROBE_CACHE.get("grouped")
        categories_present = _PROBE_CACHE.get("categories")

    if not grouped:
        info = probe_categories_from_docx(src)
        grouped = _PROBE_CACHE.get("grouped")
        if isinstance(info, dict):
            categories_present = info.get("categories")

    if not grouped:
        groups_all_tables, all_rows = read_groups_from_doc(src, progress=False)
        grouped = defaultdict(list)
        for g in groups_all_tables:
            grouped[kind_of(g["name"])].append(g)
        categories_present = [cat for cat in CATEGORY_ORDER if grouped.get(cat)]
        _PROBE_CACHE.update({
            "src": str(src),
            "grouped": grouped,
            "all_rows": all_rows,
            "categories": categories_present,
        })

    if not isinstance(grouped, defaultdict):
        tmp = defaultdict(list)
        for k, v in (grouped or {}).items():
            tmp[k] = list(v)
        grouped = tmp

    categories_present = categories_present or [cat for cat in CATEGORY_ORDER if grouped.get(cat)]
    categories_present = list(categories_present)
    if not include_support and "支撑" in categories_present:
        categories_present.remove("支撑")

    globals()["NONINTERACTIVE_MODE2_FORCE_SAME_BREAKS"] = True
    globals()["NONINTERACTIVE_MODE2_DATE_FIRST"] = (date_first or "").strip()
    globals()["NONINTERACTIVE_MODE2_DATE_SECOND"] = (date_second or "").strip()
    globals()["NONINTERACTIVE_SUPPORT_BREAKS"] = (breaks_support or "").strip()
    globals()["NONINTERACTIVE_SUPPORT_STRATEGY"] = (support_strategy or "number").lower()
    globals()["NONINTERACTIVE_NET_STRATEGY"] = (net_strategy or "number").lower()
    globals()["NONINTERACTIVE_NET_BREAKS"] = (breaks_net or "").strip()

    created_here = wb is None
    if wb is None:
        template_path = None
        for name in ("XLSX_WITH_SUPPORT_DEFAULT", "XLSX_TEMPLATE_WITH_SUPPORT", "DEFAULT_XLSX_WITH_SUPPORT"):
            if name in globals() and globals()[name]:
                template_path = Path(globals()[name])
                break
        if not template_path or not template_path.exists():
            raise FileNotFoundError("未找到 Excel 模板常量（XLSX_WITH_SUPPORT_DEFAULT / XLSX_TEMPLATE_WITH_SUPPORT / DEFAULT_XLSX_WITH_SUPPORT）。")
        wb = load_workbook_safe(template_path)

    try:
        used_pages = run_mode(
            "2",
            wb,
            categories_present=categories_present,
            grouped_preloaded=grouped,
            breaks_gz=breaks_gz or "",
            breaks_gl=breaks_gl or "",
            include_support=include_support,
        )
    finally:
        for key in (
            "NONINTERACTIVE_MODE2_FORCE_SAME_BREAKS",
            "NONINTERACTIVE_MODE2_DATE_FIRST",
            "NONINTERACTIVE_MODE2_DATE_SECOND",
            "NONINTERACTIVE_SUPPORT_BREAKS",
            "NONINTERACTIVE_SUPPORT_STRATEGY",
            "NONINTERACTIVE_NET_STRATEGY",
            "NONINTERACTIVE_NET_BREAKS",
        ):
            globals().pop(key, None)

    if created_here:
        all_rows = _PROBE_CACHE.get("all_rows")
        if all_rows:
            doc_out = build_summary_doc_with_progress(all_rows)
            set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
            save_docx_safe(doc_out, src.with_name("汇总原始记录.docx"))
        apply_meta_fixed(wb, categories_present, meta or {})
        enforce_mu_font(wb)
        cleanup_unused_sheets(wb, used_pages, bases=tuple(CATEGORY_ORDER))
        final_path = src.with_name(f"{TITLE}_报告版.xlsx")
        save_workbook_safe(wb, final_path)
        word_out = src.with_name("汇总原始记录.docx")
        return {"excel": final_path, "word": word_out}

    return {"used_pages": used_pages, "workbook": wb}

def read_groups_from_doc(path: Path, *, progress: bool = True):
    """
    从Word文档中读取并解析构件数据组，返回结构化分组数据和原始行数据。

    流程：
    1. 打开Word文档并遍历所有表格，筛选含“测点1”和“平均值”的有效数据表格
    2. 对每个有效表格提取数据行（带进度提示）
    3. 将提取的原始行数据转换为按构件名称分组的结构化数据

    结构化数据组包含构件名称和对应的测点数据（8个读数+1个平均值），适配后续Excel填充需求。

    Args:
        path: Word文档路径（Path对象）
    Returns:
        tuple: 包含两个元素的元组：
            - 构件数据组列表（list[dict]），每个元素含'name'（构件名）和'data'（数据行列表）
            - 所有原始数据行列表（list[dict]），含提取的测点值、平均值等原始信息
    """
    doc = Document(str(path))
    all_rows = []
    tables = doc.tables
    T = sum(1 for t in tables if is_data_table(t))  # noqa
    used = 0
    for tbl in tables:
        if not is_data_table(tbl):
            continue
        used += 1
        part = extract_rows_with_progress(tbl, used, T, show_progress=progress)
        if part:
            all_rows.extend(part)
    return groups_from_your_rows(all_rows), all_rows


def main():
    """命令行交互入口。"""
    print(f"{TITLE} {VERSION}")
    print("输入 help 查看模式说明；随时输入 q 返回上一步。")

    while True:
        try:
            src = prompt_path("📄 请选择原始记录 Word", WORD_SRC_DEFAULT)
        except BackStep:
            print("↩ 已返回。")
            continue
        except KeyboardInterrupt:
            print("\n已取消。")
            return
        except EOFError:
            print("\n已退出。")
            return

        try:
            probe = probe_categories_from_docx(src)
        except Exception as exc:
            print(f"❌ 识别失败：{exc}")
            continue

        categories = list((probe or {}).get("categories") or [])
        counts = (probe or {}).get("counts") or {}
        if categories:
            print("📊 识别：" + "、".join(f"{cat} {counts.get(cat, 0)}" for cat in categories))
        else:
            print("⚠️ 未识别到可用构件。")

        all_rows = _PROBE_CACHE.get("all_rows")
        if all_rows:
            try:
                doc_out = build_summary_doc_with_progress(all_rows)
                set_doc_font_progress(doc_out, DEFAULT_FONT_PT)
                out_docx = Path(src).with_name("汇总原始记录.docx")
                print("💾 正在保存汇总 Word …")
                save_docx_safe(doc_out, out_docx)
                print(f"✅ 汇总 Word 已保存：{out_docx}")
            except Exception as exc:
                print(f"⚠️ 汇总 Word 保存失败：{exc}")

        try:
            proj = ask("工程名称（回车跳过，输入 q 返回）：")
            order = ask("委托编号（回车跳过，输入 q 返回）：")
        except BackStep:
            print("↩ 返回文件选择。")
            continue

        meta = {}
        if proj:
            meta["proj"] = proj
        if order:
            meta["order"] = order

        grouped_cached = _PROBE_CACHE.get("grouped") or defaultdict(list)
        try:
            run_with_mode(Path(src), grouped_cached, categories or None, meta)
        except BackStep:
            print("↩ 返回模式选择。")
            continue
        except Exception as exc:
            print(f"❌ 出表失败：{exc}")
            continue

        try:
            again = ask("是否继续处理其他文件？（y=继续 / 其它=退出）：", lower=True)
        except BackStep:
            break
        if again != "y":
            break


if __name__ == "__main__":
    main()

                                                                                                         # v 1.0.1
